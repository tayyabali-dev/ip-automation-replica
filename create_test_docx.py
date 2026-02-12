#!/usr/bin/env python3
"""
Create a proper DOCX file with patent application content for testing.
"""

from docx import Document
from docx.shared import Inches

def create_test_docx():
    """Create a test DOCX file with patent application data."""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Patent Application Data Sheet', 0)
    
    # Add application info
    doc.add_heading('Application Information', level=1)
    p = doc.add_paragraph()
    p.add_run('Title of Invention: ').bold = True
    p.add_run('Advanced AI-Powered Document Processing System')
    
    p = doc.add_paragraph()
    p.add_run('Application Number: ').bold = True
    p.add_run('17/123,456')
    
    p = doc.add_paragraph()
    p.add_run('Application Type: ').bold = True
    p.add_run('Utility')
    
    p = doc.add_paragraph()
    p.add_run('Entity Status: ').bold = True
    p.add_run('Small Entity')
    
    p = doc.add_paragraph()
    p.add_run('Total Drawing Sheets: ').bold = True
    p.add_run('5')
    
    p = doc.add_paragraph()
    p.add_run('Suggested Representative Figure: ').bold = True
    p.add_run('1')
    
    # Add inventor information
    doc.add_heading('Inventor Information', level=1)
    
    # Create a table for inventor info
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Address'
    hdr_cells[2].text = 'City, State, ZIP'
    hdr_cells[3].text = 'Citizenship'
    
    # Add inventor 1
    row_cells = table.add_row().cells
    row_cells[0].text = 'John Michael Smith'
    row_cells[1].text = '123 Innovation Drive'
    row_cells[2].text = 'Palo Alto, CA, 94301'
    row_cells[3].text = 'United States'
    
    # Add inventor 2
    row_cells = table.add_row().cells
    row_cells[0].text = 'Sarah Elizabeth Johnson'
    row_cells[1].text = '456 Technology Lane'
    row_cells[2].text = 'Cambridge, MA, 02139'
    row_cells[3].text = 'United States'
    
    # Add applicant information
    doc.add_heading('Applicant Information', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Applicant 1: ').bold = True
    p.add_run('TechCorp Innovations LLC')
    
    p = doc.add_paragraph()
    p.add_run('Address: ').bold = True
    p.add_run('789 Corporate Blvd, Suite 100')
    
    p = doc.add_paragraph()
    p.add_run('City, State, ZIP: ').bold = True
    p.add_run('San Francisco, CA, 94105')
    
    p = doc.add_paragraph()
    p.add_run('Country: ').bold = True
    p.add_run('United States')
    
    # Add correspondence address
    doc.add_heading('Correspondence Address', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Law Firm: ').bold = True
    p.add_run('Blakely, Sokoloff, Taylor & Zafman LLP')
    
    p = doc.add_paragraph()
    p.add_run('Address: ').bold = True
    p.add_run('1279 Oakmead Parkway')
    
    p = doc.add_paragraph()
    p.add_run('City, State, ZIP: ').bold = True
    p.add_run('Sunnyvale, CA, 94085')
    
    p = doc.add_paragraph()
    p.add_run('Country: ').bold = True
    p.add_run('United States')
    
    p = doc.add_paragraph()
    p.add_run('Phone: ').bold = True
    p.add_run('(408) 720-8000')
    
    p = doc.add_paragraph()
    p.add_run('Email: ').bold = True
    p.add_run('patents@bstlaw.com')
    
    p = doc.add_paragraph()
    p.add_run('Customer Number: ').bold = True
    p.add_run('12345')
    
    # Add abstract section
    doc.add_heading('Abstract', level=1)
    abstract_text = """
    An advanced AI-powered document processing system that automatically extracts 
    patent application data from various document formats including PDF and DOCX files. 
    The system uses machine learning algorithms to identify and parse inventor information, 
    applicant details, correspondence addresses, and other critical patent metadata 
    with high accuracy and efficiency.
    """
    doc.add_paragraph(abstract_text)
    
    # Save the document
    doc.save('test_data/patent_application_test.docx')
    print("✅ Created test_data/patent_application_test.docx")

if __name__ == "__main__":
    create_test_docx()