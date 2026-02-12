import io
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

from app.models.office_action import (
    OfficeActionExtractedData,
    Rejection,
    ClaimStatus,
    Objection,
    ExaminerStatement,
    PriorArtReference
)

logger = logging.getLogger(__name__)


class ReportGenerator:
    def generate_office_action_report(self, data: Dict[str, Any]) -> io.BytesIO:
        """
        Generates a Word document report for a Patent Office Action.
        """
        # Convert dict to model if necessary
        if isinstance(data, dict):
            try:
                oa_data = OfficeActionExtractedData(**data)
            except Exception as e:
                logger.error(f"Failed to parse data into model: {e}")
                raise ValueError("Invalid Office Action Data format")
        else:
            oa_data = data

        document = Document()

        # --- TITLE ---
        title = document.add_heading('Office Action Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- SECTION 1: APPLICATION SUMMARY ---
        document.add_heading('1. Application Summary', level=1)

        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Details'

        # Data Rows - using empty string for missing values instead of N/A
        def format_value(val):
            if val is None:
                return ""
            return str(val)

        header_info = [
            ("Application Number", oa_data.header.application_number),
            ("Filing Date", oa_data.header.filing_date),
            ("Office Action Date", oa_data.header.office_action_date),
            ("Type", oa_data.header.office_action_type),
            ("Examiner", oa_data.header.examiner_name),
            ("Examiner Type", oa_data.header.examiner_type),
            ("Examiner Phone", oa_data.header.examiner_phone),
            ("Examiner Email", oa_data.header.examiner_email),
            ("Art Unit", oa_data.header.art_unit),
            ("Response Deadline", oa_data.header.response_deadline),
            ("Confirmation No.", oa_data.header.confirmation_number),
            ("Attorney Docket No.", oa_data.header.attorney_docket_number),
            ("Customer Number", oa_data.header.customer_number),
            ("Title of Invention", oa_data.header.title_of_invention),
            ("First Named Inventor", oa_data.header.first_named_inventor),
            ("Applicant Name", oa_data.header.applicant_name),
        ]

        for key, value in header_info:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = format_value(value)

        document.add_paragraph()  # Spacer

        # --- SECTION 1A: FOREIGN PRIORITY (if present) ---
        if oa_data.header.foreign_priority:
            document.add_heading('1a. Foreign Priority Claims', level=2)
            fp_table = document.add_table(rows=1, cols=3)
            fp_table.style = 'Table Grid'
            hdr = fp_table.rows[0].cells
            hdr[0].text = 'Country'
            hdr[1].text = 'Application Number'
            hdr[2].text = 'Filing Date'

            for fp in oa_data.header.foreign_priority:
                row = fp_table.add_row().cells
                row[0].text = format_value(fp.country)
                row[1].text = format_value(fp.application_number)
                row[2].text = format_value(fp.filing_date)

            document.add_paragraph()

        # --- SECTION 1B: PARENT APPLICATIONS (if present) ---
        if oa_data.header.parent_applications:
            document.add_heading('1b. Parent Application Data (Continuity)', level=2)
            pa_table = document.add_table(rows=1, cols=4)
            pa_table.style = 'Table Grid'
            hdr = pa_table.rows[0].cells
            hdr[0].text = 'Parent App. No.'
            hdr[1].text = 'Filing Date'
            hdr[2].text = 'Relationship'
            hdr[3].text = 'Status'

            for pa in oa_data.header.parent_applications:
                row = pa_table.add_row().cells
                row[0].text = format_value(pa.parent_application_number)
                row[1].text = format_value(pa.parent_filing_date)
                row[2].text = format_value(pa.relationship_type)
                row[3].text = format_value(pa.status)

            document.add_paragraph()

        # --- SECTION 2: CLAIMS STATUS OVERVIEW ---
        document.add_heading('2. Claims Status Overview', level=1)

        if oa_data.claims_status:
            # Determine if we have claim text to show
            has_claim_text = any(c.claim_text for c in oa_data.claims_status)

            if has_claim_text:
                claims_table = document.add_table(rows=1, cols=4)
            else:
                claims_table = document.add_table(rows=1, cols=4)

            claims_table.style = 'Table Grid'
            hdr_cells = claims_table.rows[0].cells
            hdr_cells[0].text = 'Claim No.'
            hdr_cells[1].text = 'Status'
            hdr_cells[2].text = 'Type'
            hdr_cells[3].text = 'Parent'

            for claim in oa_data.claims_status:
                row_cells = claims_table.add_row().cells
                row_cells[0].text = claim.claim_number
                row_cells[1].text = claim.status
                row_cells[2].text = claim.dependency_type
                row_cells[3].text = format_value(claim.parent_claim)

            # If we have claim text, add a separate section
            if has_claim_text:
                document.add_paragraph()
                document.add_heading('2a. Claim Text (Extracted)', level=2)
                for claim in oa_data.claims_status:
                    if claim.claim_text:
                        p = document.add_paragraph()
                        p.add_run(f"Claim {claim.claim_number}: ").bold = True
                        p.add_run(claim.claim_text)
                        document.add_paragraph()
        else:
            document.add_paragraph("No specific claim status information extracted.")

        document.add_paragraph()

        # --- SECTION 3: DETAILED REJECTION ANALYSIS ---
        document.add_heading('3. Detailed Rejection Analysis', level=1)

        if not oa_data.rejections:
            document.add_paragraph("No rejections found in this Office Action.")

        for i, rejection in enumerate(oa_data.rejections):
            # Rejection Header with normalized type if available
            rejection_label = rejection.rejection_type
            if rejection.rejection_type_normalized:
                rejection_label = f"{rejection.rejection_type_normalized}"

            heading = document.add_heading(f'Rejection #{i+1}: {rejection_label}', level=2)

            # AIA indicator
            if rejection.is_aia is not None:
                p = document.add_paragraph()
                p.add_run("Statute Version: ").bold = True
                p.add_run("AIA" if rejection.is_aia else "Pre-AIA")

            # Details
            p = document.add_paragraph()
            p.add_run("Claims Affected: ").bold = True
            p.add_run(", ".join(rejection.affected_claims) if rejection.affected_claims else "Not specified")

            p = document.add_paragraph()
            p.add_run("Statutory Basis: ").bold = True
            p.add_run(rejection.statutory_basis or "Not specified")

            # Prior Art
            document.add_heading('Prior Art Cited:', level=3)
            if rejection.cited_prior_art:
                for art in rejection.cited_prior_art:
                    p = document.add_paragraph(style='List Bullet')
                    txt = f"{art.identifier}"
                    if art.short_name:
                        txt = f"{art.short_name} ({art.identifier})"
                    if art.title:
                        txt += f" - {art.title}"
                    if art.relevant_sections:
                        txt += f" [{art.relevant_sections}]"
                    p.add_run(txt)
            else:
                document.add_paragraph("No specific prior art references cited for this rejection.")

            # Section 103 Combinations (if applicable)
            if rejection.prior_art_combinations:
                document.add_heading('Prior Art Combinations (103):', level=3)
                for j, combo in enumerate(rejection.prior_art_combinations):
                    p = document.add_paragraph()
                    p.add_run(f"Combination {j+1}: ").bold = True

                    p = document.add_paragraph(style='List Bullet')
                    p.add_run("Primary: ").bold = True
                    p.add_run(combo.primary_reference_id)

                    if combo.secondary_reference_ids:
                        p = document.add_paragraph(style='List Bullet')
                        p.add_run("Secondary: ").bold = True
                        p.add_run(", ".join(combo.secondary_reference_ids))

                    if combo.motivation_to_combine:
                        p = document.add_paragraph(style='List Bullet')
                        p.add_run("Motivation to Combine: ").bold = True
                        p.add_run(combo.motivation_to_combine)

            # Reasoning
            document.add_heading("Examiner's Reasoning:", level=3)
            document.add_paragraph(rejection.examiner_reasoning)

            # Separator
            document.add_paragraph('_' * 40)

        # --- SECTION 4: OBJECTIONS ---
        if oa_data.objections:
            document.add_heading('4. Objections', level=1)
            for i, obj in enumerate(oa_data.objections):
                obj_type = ""
                if obj.objection_type:
                    obj_type = f" ({obj.objection_type})"
                document.add_heading(f"Objection #{i+1}: {obj.objected_item}{obj_type}", level=2)
                p = document.add_paragraph()
                p.add_run("Reason: ").bold = True
                p.add_run(obj.reason)
                if obj.corrective_action:
                    p = document.add_paragraph()
                    p.add_run("Required Action: ").bold = True
                    p.add_run(obj.corrective_action)

        # --- SECTION 5: EXAMINER STATEMENTS ---
        if oa_data.other_statements:
            document.add_heading('5. Examiner Comments & Allowable Subject Matter', level=1)
            for stmt in oa_data.other_statements:
                document.add_heading(stmt.statement_type, level=2)
                document.add_paragraph(stmt.content)

        # --- SECTION 6: CONSOLIDATED PRIOR ART REFERENCES ---
        if oa_data.all_references:
            document.add_heading('6. All Prior Art References', level=1)
            ref_table = document.add_table(rows=1, cols=5)
            ref_table.style = 'Table Grid'
            hdr = ref_table.rows[0].cells
            hdr[0].text = 'ID'
            hdr[1].text = 'Reference'
            hdr[2].text = 'Type'
            hdr[3].text = 'Cited Sections'
            hdr[4].text = 'Used In Rejections'

            for ref in oa_data.all_references:
                row = ref_table.add_row().cells
                row[0].text = format_value(ref.reference_id)
                ref_name = ref.identifier
                if ref.short_name:
                    ref_name = f"{ref.short_name} ({ref.identifier})"
                row[1].text = ref_name
                row[2].text = ref.reference_type
                row[3].text = format_value(ref.relevant_sections)
                if ref.used_in_rejection_indices:
                    row[4].text = ", ".join([f"#{idx+1}" for idx in ref.used_in_rejection_indices])
                else:
                    row[4].text = ""

            document.add_paragraph()

        # --- SECTION 7: DEADLINE CALCULATION ---
        document.add_heading('7. Response Deadlines', level=1)

        if oa_data.deadline_calculation:
            dc = oa_data.deadline_calculation

            p = document.add_paragraph()
            p.add_run("Mailing Date: ").bold = True
            p.add_run(dc.mailing_date)

            p = document.add_paragraph()
            p.add_run("Shortened Statutory Period: ").bold = True
            p.add_run(f"{dc.shortened_statutory_period} months")

            p = document.add_paragraph()
            p.add_run("Statutory Deadline (no extension): ").bold = True
            p.add_run(dc.statutory_deadline)

            p = document.add_paragraph()
            p.add_run("Maximum Deadline (with extensions): ").bold = True
            p.add_run(dc.maximum_deadline)

            if dc.is_final_action:
                p = document.add_paragraph()
                p.add_run("NOTE: This is a Final Office Action. Extension options may be limited.").bold = True

            # Deadline tiers table
            if dc.tiers:
                document.add_paragraph()
                document.add_heading('Deadline Tiers with Extension Fees:', level=2)

                tier_table = document.add_table(rows=1, cols=5)
                tier_table.style = 'Table Grid'
                hdr = tier_table.rows[0].cells
                hdr[0].text = 'Deadline'
                hdr[1].text = 'Extension'
                hdr[2].text = 'Micro ($)'
                hdr[3].text = 'Small ($)'
                hdr[4].text = 'Large ($)'

                for tier in dc.tiers:
                    row = tier_table.add_row().cells
                    deadline_text = tier.deadline_date
                    if tier.is_past:
                        deadline_text += " (PAST)"
                    row[0].text = deadline_text
                    row[1].text = f"{tier.months_extension} month(s)" if tier.months_extension > 0 else "None"
                    row[2].text = f"${tier.extension_fee_micro:,}"
                    row[3].text = f"${tier.extension_fee_small:,}"
                    row[4].text = f"${tier.extension_fee_large:,}"

            # Notes
            if dc.notes:
                document.add_paragraph()
                p = document.add_paragraph()
                p.add_run("Notes:").bold = True
                for note in dc.notes:
                    document.add_paragraph(f"- {note}", style='List Bullet')
        else:
            p = document.add_paragraph()
            p.add_run("Stated Deadline: ").bold = True
            p.add_run(oa_data.header.response_deadline or "Please verify manually based on mailing date.")
            document.add_paragraph()
            document.add_paragraph(
                "Note: Use the 'Calculate Deadlines' feature for detailed deadline tiers with extension fees."
            )

        # Save to IO stream
        output = io.BytesIO()
        document.save(output)
        output.seek(0)

        return output


report_generator = ReportGenerator()
