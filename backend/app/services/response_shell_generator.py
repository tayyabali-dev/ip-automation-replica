"""
Response Shell Generator Service for USPTO Office Actions.

Generates a USPTO Response document template (DOCX) with proper legal formatting,
including placeholders for attorney arguments.
"""
import io
import logging
from datetime import date
from typing import Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.office_action import (
    OfficeActionExtractedData,
    Rejection,
    Objection,
    ClaimStatus
)

logger = logging.getLogger(__name__)


class FirmInfo:
    """Firm and attorney information for signature block."""
    def __init__(
        self,
        firm_name: Optional[str] = None,
        attorney_name: Optional[str] = None,
        attorney_reg_number: Optional[str] = None,
        firm_address: Optional[str] = None,
        firm_phone: Optional[str] = None,
        firm_email: Optional[str] = None
    ):
        self.firm_name = firm_name or "[FIRM NAME]"
        self.attorney_name = attorney_name or "[ATTORNEY NAME]"
        self.attorney_reg_number = attorney_reg_number or "[REG. NO.]"
        self.firm_address = firm_address or "[FIRM ADDRESS]"
        self.firm_phone = firm_phone or "[PHONE]"
        self.firm_email = firm_email or "[EMAIL]"


class ResponseShellGenerator:
    """
    Generates a USPTO Office Action Response template (DOCX).

    Includes:
    - Header block with application info
    - Amendment section (for claim amendments)
    - Remarks section with rejection-specific placeholders
    - Conclusion
    - Signature block
    """

    def __init__(self):
        # Standard USPTO fonts and sizes
        self.body_font = "Times New Roman"
        self.body_size = Pt(12)
        self.heading_font = "Times New Roman"
        self.heading_size = Pt(14)

    def _set_default_styles(self, doc: Document):
        """Set default document styles for USPTO compliance."""
        # Set default paragraph style
        style = doc.styles['Normal']
        font = style.font
        font.name = self.body_font
        font.size = self.body_size

        # Set document margins (USPTO: 1 inch minimum on all sides)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_centered_bold_text(self, doc: Document, text: str, size: Pt = None):
        """Add centered, bold text."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        if size:
            run.font.size = size
        return p

    def _add_placeholder(self, doc: Document, text: str):
        """Add a placeholder text that's easily identifiable."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        # Highlight would be nice but keeping it simple
        return p

    def generate_response_shell(
        self,
        data: OfficeActionExtractedData,
        firm_info: Optional[FirmInfo] = None,
        include_claim_amendments: bool = True,
        include_specification_amendments: bool = False
    ) -> io.BytesIO:
        """
        Generate a response shell document.

        Args:
            data: Extracted Office Action data
            firm_info: Firm and attorney information for signature block
            include_claim_amendments: Include claims amendment section
            include_specification_amendments: Include specification amendment section

        Returns:
            BytesIO containing the DOCX document
        """
        if firm_info is None:
            firm_info = FirmInfo()

        doc = Document()
        self._set_default_styles(doc)

        # 1. Header Block
        self._add_header_block(doc, data)

        # 2. Commissioner Address
        self._add_commissioner_address(doc, data)

        # 3. Amendment Section (if claims are being amended)
        if include_claim_amendments:
            self._add_claims_amendment_section(doc, data)

        if include_specification_amendments:
            self._add_specification_amendment_section(doc)

        # 4. Remarks Section
        self._add_remarks_section(doc, data)

        # 5. Conclusion
        self._add_conclusion(doc)

        # 6. Signature Block
        self._add_signature_block(doc, firm_info)

        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def _add_header_block(self, doc: Document, data: OfficeActionExtractedData):
        """Add the standard USPTO response header."""
        # PATENT designation
        self._add_centered_bold_text(doc, "PATENT", self.heading_size)
        doc.add_paragraph()

        # Main header
        self._add_centered_bold_text(doc, "IN THE UNITED STATES PATENT AND TRADEMARK OFFICE")
        doc.add_paragraph()

        # Two-column application info table
        table = doc.add_table(rows=7, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(3)
        for cell in table.columns[1].cells:
            cell.width = Inches(3)

        # Populate table with application info
        info_rows = [
            ("In re Application of:", data.header.first_named_inventor or "[INVENTOR NAME]"),
            ("Serial No.:", data.header.application_number or "[APPLICATION NUMBER]"),
            ("Filed:", data.header.filing_date or "[FILING DATE]"),
            ("Confirmation No.:", data.header.confirmation_number or "[CONFIRMATION NO.]"),
            ("Examiner:", data.header.examiner_name or "[EXAMINER NAME]"),
            ("Art Unit:", data.header.art_unit or "[ART UNIT]"),
            ("Attorney Docket No.:", data.header.attorney_docket_number or "[DOCKET NO.]"),
        ]

        for i, (label, value) in enumerate(info_rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value) if value else ""

        doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.add_run("For: ").bold = True
        p.add_run(data.header.title_of_invention or "[TITLE OF INVENTION]")

        doc.add_paragraph()

        # Response header
        self._add_centered_bold_text(doc, "RESPONSE TO OFFICE ACTION", self.heading_size)
        doc.add_paragraph()

    def _add_commissioner_address(self, doc: Document, data: OfficeActionExtractedData):
        """Add the Commissioner address block."""
        p = doc.add_paragraph()
        p.add_run("Commissioner for Patents")
        doc.add_paragraph("P.O. Box 1450")
        doc.add_paragraph("Alexandria, VA 22313-1450")
        doc.add_paragraph()

        # Reference to Office Action
        p = doc.add_paragraph()
        oa_date = data.header.office_action_date or "[OFFICE ACTION DATE]"
        p.add_run(f"Dear Sir/Madam:")
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(
            f"In response to the Office Action mailed {oa_date}, "
            f"please consider the following amendments and remarks."
        )
        doc.add_paragraph()

    def _add_claims_amendment_section(self, doc: Document, data: OfficeActionExtractedData):
        """Add claims amendment section with placeholders."""
        # Section header
        heading = doc.add_heading("AMENDMENTS TO THE CLAIMS", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Boilerplate text
        p = doc.add_paragraph()
        p.add_run(
            "This listing of claims will replace all prior versions and listings "
            "of claims in the application:"
        )
        doc.add_paragraph()

        # Listing of claims header
        p = doc.add_paragraph()
        p.add_run("Listing of Claims:").bold = True
        doc.add_paragraph()

        # Add placeholder for each claim
        for claim in data.claims_status:
            status_marker = self._get_claim_status_marker(claim)

            p = doc.add_paragraph()
            p.add_run(f"{claim.claim_number}. ").bold = True
            p.add_run(f"({status_marker}) ")

            # Add claim text if available, otherwise placeholder
            if claim.claim_text:
                p.add_run(claim.claim_text)
            else:
                run = p.add_run("[INSERT CLAIM TEXT]")
                run.italic = True

            doc.add_paragraph()

        doc.add_paragraph()

    def _get_claim_status_marker(self, claim: ClaimStatus) -> str:
        """Get the USPTO claim status marker."""
        status_lower = claim.status.lower()
        if "amend" in status_lower or "rejected" in status_lower:
            return "Currently Amended"
        elif "allow" in status_lower:
            return "Previously Presented"
        elif "cancel" in status_lower:
            return "Cancelled"
        elif "withdraw" in status_lower:
            return "Withdrawn"
        elif "new" in status_lower:
            return "New"
        else:
            return "Original"

    def _add_specification_amendment_section(self, doc: Document):
        """Add specification amendment section with placeholders."""
        heading = doc.add_heading("AMENDMENTS TO THE SPECIFICATION", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        self._add_placeholder(doc, "[INSERT SPECIFICATION AMENDMENTS IF NEEDED]")
        doc.add_paragraph()

    def _add_remarks_section(self, doc: Document, data: OfficeActionExtractedData):
        """Add remarks section with rejection-specific placeholders."""
        heading = doc.add_heading("REMARKS", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Intro paragraph
        p = doc.add_paragraph()
        p.add_run(
            "Applicant respectfully submits the following remarks in support of "
            "patentability of the claims."
        )
        doc.add_paragraph()

        # Handle each rejection type
        rejection_count = 0
        for i, rejection in enumerate(data.rejections):
            rejection_count += 1
            self._add_rejection_response(doc, rejection, rejection_count)

        # Handle objections
        objection_count = 0
        for objection in data.objections:
            objection_count += 1
            self._add_objection_response(doc, objection, objection_count)

        # If no rejections or objections
        if rejection_count == 0 and objection_count == 0:
            self._add_placeholder(doc, "[NO REJECTIONS OR OBJECTIONS IDENTIFIED - REVIEW OFFICE ACTION]")

    def _add_rejection_response(self, doc: Document, rejection: Rejection, index: int):
        """Add response template for a rejection."""
        # Rejection heading
        statutory_basis = rejection.statutory_basis or f"35 U.S.C. {rejection.rejection_type}"
        doc.add_heading(f"Response to {statutory_basis} Rejection", level=2)

        # Claims affected
        claims_str = ", ".join(rejection.affected_claims) if rejection.affected_claims else "[CLAIMS]"
        p = doc.add_paragraph()
        p.add_run(f"Claims {claims_str} ").bold = True
        p.add_run(f"stand rejected under {statutory_basis}.")
        doc.add_paragraph()

        # Add specific response template based on rejection type
        rejection_type = rejection.rejection_type.lower()
        if "103" in rejection_type:
            self._add_103_response_template(doc, rejection)
        elif "102" in rejection_type:
            self._add_102_response_template(doc, rejection)
        elif "112" in rejection_type:
            self._add_112_response_template(doc, rejection)
        elif "101" in rejection_type:
            self._add_101_response_template(doc, rejection)
        elif "double" in rejection_type.lower():
            self._add_double_patenting_response_template(doc, rejection)
        else:
            self._add_generic_response_template(doc, rejection)

        doc.add_paragraph()

    def _add_103_response_template(self, doc: Document, rejection: Rejection):
        """Add 103 obviousness response template."""
        p = doc.add_paragraph()
        p.add_run(
            "Applicant respectfully traverses the rejection of claims under 35 U.S.C. 103."
        )
        doc.add_paragraph()

        # List cited references
        if rejection.cited_prior_art:
            p = doc.add_paragraph()
            refs = []
            for art in rejection.cited_prior_art:
                ref_name = art.short_name or art.identifier
                refs.append(ref_name)
            p.add_run(f"The Examiner relies on {', '.join(refs)}.")
            doc.add_paragraph()

        # Show combination if available
        if rejection.prior_art_combinations:
            for combo in rejection.prior_art_combinations:
                p = doc.add_paragraph()
                p.add_run("Combination: ").bold = True
                secondary_refs = ", ".join(combo.secondary_reference_ids) if combo.secondary_reference_ids else "none"
                p.add_run(f"Primary: {combo.primary_reference_id}, Secondary: {secondary_refs}")

                if combo.motivation_to_combine:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run("Examiner's Motivation: ").bold = True
                    p.add_run(combo.motivation_to_combine)

        doc.add_paragraph()

        # Argument placeholders - comprehensive set for §103 obviousness
        self._add_placeholder(doc, "[ARGUMENT 1: The cited references, alone or in combination, fail to teach or suggest _____________]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[ARGUMENT 2: One of ordinary skill in the art would not have been motivated to combine the references because _____________]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[ARGUMENT 3: The combination would not render the claimed invention obvious because _____________]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[ARGUMENT 4: The cited art teaches away from the claimed invention because _____________]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[ARGUMENT 5: The claimed invention achieves unexpected results including _____________]")

    def _add_102_response_template(self, doc: Document, rejection: Rejection):
        """Add 102 novelty response template."""
        p = doc.add_paragraph()
        p.add_run(
            "Applicant respectfully traverses the rejection of claims under 35 U.S.C. 102."
        )
        doc.add_paragraph()

        # List cited reference
        if rejection.cited_prior_art:
            p = doc.add_paragraph()
            ref = rejection.cited_prior_art[0]
            ref_name = ref.short_name or ref.identifier
            p.add_run(f"The Examiner cites {ref_name} as anticipating the claimed invention.")
            doc.add_paragraph()

        # Argument placeholders
        self._add_placeholder(doc, "[ARGUMENT: The reference does not disclose the claim limitation of _____________]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[ARGUMENT: The reference teaches away from the claimed invention because _____________]")

    def _add_112_response_template(self, doc: Document, rejection: Rejection):
        """Add 112 response template."""
        p = doc.add_paragraph()
        p.add_run(
            "Applicant respectfully traverses the rejection of claims under 35 U.S.C. 112."
        )
        doc.add_paragraph()

        # Determine subsection
        rejection_type = rejection.rejection_type_normalized or rejection.rejection_type
        if "112(a)" in str(rejection_type) or "first" in str(rejection_type).lower():
            doc.add_paragraph(
                "Regarding the written description/enablement requirement:"
            )
            self._add_placeholder(doc, "[ARGUMENT: The specification provides adequate support at paragraphs/pages _____________]")
        elif "112(b)" in str(rejection_type) or "second" in str(rejection_type).lower():
            doc.add_paragraph(
                "Regarding the definiteness requirement:"
            )
            self._add_placeholder(doc, "[ARGUMENT: The claim language is clear because _____________]")
        else:
            self._add_placeholder(doc, "[ARGUMENT: Address the specific 112 issue identified by the Examiner]")

    def _add_101_response_template(self, doc: Document, rejection: Rejection):
        """Add 101 eligibility response template."""
        p = doc.add_paragraph()
        p.add_run(
            "Applicant respectfully traverses the rejection of claims under 35 U.S.C. 101."
        )
        doc.add_paragraph()

        # Alice/Mayo framework
        doc.add_paragraph("Under the Alice/Mayo framework:")
        doc.add_paragraph()

        self._add_placeholder(doc, "[STEP 2A PRONG 1: The claims are not directed to an abstract idea because _____________]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[STEP 2A PRONG 2: Even if directed to an exception, the claims integrate it into a practical application because _____________]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[STEP 2B: The claims recite significantly more than the abstract idea because _____________]")

    def _add_double_patenting_response_template(self, doc: Document, rejection: Rejection):
        """Add double patenting response template."""
        p = doc.add_paragraph()
        p.add_run(
            "Applicant respectfully addresses the double patenting rejection."
        )
        doc.add_paragraph()

        self._add_placeholder(doc, "[OPTION 1: File a terminal disclaimer to obviate the rejection]")
        doc.add_paragraph()
        self._add_placeholder(doc, "[OPTION 2: Argue that the claims are patentably distinct because _____________]")

    def _add_generic_response_template(self, doc: Document, rejection: Rejection):
        """Add generic response template for other rejection types."""
        p = doc.add_paragraph()
        p.add_run(
            "Applicant respectfully traverses this rejection."
        )
        doc.add_paragraph()

        self._add_placeholder(doc, "[INSERT ARGUMENTS ADDRESSING THE EXAMINER'S REJECTION]")

    def _add_objection_response(self, doc: Document, objection: Objection, index: int):
        """Add response template for an objection."""
        doc.add_heading(f"Response to {objection.objected_item} Objection", level=2)

        p = doc.add_paragraph()
        p.add_run("Objection: ").bold = True
        p.add_run(objection.reason)
        doc.add_paragraph()

        if objection.corrective_action:
            p = doc.add_paragraph()
            p.add_run("Suggested Correction: ").bold = True
            p.add_run(objection.corrective_action)
            doc.add_paragraph()

        self._add_placeholder(doc, "[INSERT RESPONSE TO OBJECTION OR CONFIRM CORRECTION MADE]")
        doc.add_paragraph()

    def _add_conclusion(self, doc: Document):
        """Add conclusion section."""
        heading = doc.add_heading("CONCLUSION", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(
            "In view of the foregoing amendments and remarks, Applicant respectfully "
            "submits that all pending claims are in condition for allowance. "
            "Favorable reconsideration and prompt allowance of this application is respectfully requested."
        )
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(
            "If any extensions of time are needed to enter this response, "
            "such extensions are hereby petitioned under 37 C.F.R. 1.136(a), "
            "and any fees required therefor are hereby authorized to be charged "
            "to Deposit Account No. [DEPOSIT ACCOUNT] or to be charged to any credit card "
            "on file with the USPTO."
        )
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(
            "Should the Examiner have any questions or wish to discuss this response, "
            "the Examiner is invited to contact the undersigned at the telephone number below."
        )
        doc.add_paragraph()

    def _add_signature_block(self, doc: Document, firm_info: FirmInfo):
        """Add signature block."""
        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("Respectfully submitted,")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        # Signature line
        p = doc.add_paragraph()
        p.add_run("_" * 40)
        doc.add_paragraph()

        # Attorney info
        p = doc.add_paragraph()
        p.add_run(firm_info.attorney_name)
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"Registration No. {firm_info.attorney_reg_number}")
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(firm_info.firm_name)
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(firm_info.firm_address)
        doc.add_paragraph()

        if firm_info.firm_phone != "[PHONE]":
            p = doc.add_paragraph()
            p.add_run(f"Tel: {firm_info.firm_phone}")

        if firm_info.firm_email != "[EMAIL]":
            p = doc.add_paragraph()
            p.add_run(f"Email: {firm_info.firm_email}")

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"Date: {date.today().strftime('%B %d, %Y')}")


# Singleton instance
response_shell_generator = ResponseShellGenerator()
