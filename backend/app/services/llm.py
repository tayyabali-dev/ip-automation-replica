from google import genai
from google.genai import types
from google.api_core.exceptions import ResourceExhausted
from fastapi import HTTPException, status
from app.core.config import settings
from app.models.patent_application import PatentApplicationMetadata
# from app.models.extraction import ExtractionMetadata, ExtractionResult, ConfidenceLevel, DocumentQuality
from app.models.extraction import ExtractionResult
import logging
import json
import time
import re
import os
import asyncio
import io
import random
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple, Callable, Awaitable, Union, IO
from pypdf import PdfReader, PdfWriter
# Configure logging
logger = logging.getLogger(__name__)

try:
    import fitz  # PyMuPDF
    logger.info(f"PyMuPDF (fitz) imported successfully. Version: {fitz.__version__}")
except ImportError:
    fitz = None
    logger.warning("PyMuPDF (fitz) could not be imported. Image-based extraction will be unavailable.")


# ══════════════════════════════════════════════════════════════════════════════
# TEXT CLEANING UTILITIES - Fix for Fragmented PDF Text Extraction
# ══════════════════════════════════════════════════════════════════════════════

def clean_fragmented_text(text: str) -> str:
    """
    Fix PDFs with fragmented text where each word is on its own line.
    
    Some PDF creators (especially from presentations or certain export tools)
    store each word as a separate positioned text object, resulting in:
        UNITED\n \nSTATES\n \nPATENT
    
    This function normalizes such text to:
        UNITED STATES PATENT
    
    Args:
        text: Raw extracted text from PDF
        
    Returns:
        Cleaned text with proper spacing
    """
    if not text:
        return text
    
    # Pattern 1: newline + space + newline (most common fragmentation)
    # "WORD\n \nWORD" -> "WORD WORD"
    text = text.replace('\n \n', ' ')
    
    # Pattern 2: newline + multiple spaces + newline
    # "WORD\n   \nWORD" -> "WORD WORD"
    text = re.sub(r'\n\s+\n', ' ', text)
    
    # Pattern 3: Multiple consecutive newlines (preserve paragraph breaks as double newline)
    # "sentence.\n\n\n\nNew paragraph" -> "sentence.\n\nNew paragraph"
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Pattern 4: Single newlines that break words mid-sentence
    # But preserve intentional line breaks (after periods, colons, etc.)
    # This is tricky - only collapse newlines NOT preceded by sentence-ending punctuation
    # "relates to VR\nperipherals" -> "relates to VR peripherals"
    # "Technical Field\n[0001]" -> keep as is (section break)
    text = re.sub(r'(?<![.!?:\]\)])\n(?=[a-z])', ' ', text)
    
    # Pattern 5: Collapse multiple spaces into single space
    text = re.sub(r'  +', ' ', text)
    
    # Pattern 6: Fix spaces before punctuation
    # "VANCOUVER ( CA )" -> "VANCOUVER (CA)"
    text = re.sub(r'\s+([,.:;!?\)])', r'\1', text)
    text = re.sub(r'\(\s+', '(', text)
    
    # Pattern 7: Trim whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Pattern 8: Remove empty lines that aren't paragraph breaks
    text = re.sub(r'\n\n\n+', '\n\n', text)
    
    return text.strip()


def _parse_claim_numbers(claim_text: str) -> list:
    """
    Parse claim numbers from text like "1, 8, 12 and 16-18" into ["1", "8", "12", "16", "17", "18"]
    """
    import re
    claims = []
    
    # Remove "and" and normalize
    claim_text = claim_text.replace(" and ", ", ").replace("and", ",")
    
    # Split by comma
    parts = claim_text.split(",")
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
        
        # Check if it's a range (e.g., "16-18")
        range_match = re.match(r'(\d+)\s*[-–]\s*(\d+)', part)
        if range_match:
            start, end = int(range_match.group(1)), int(range_match.group(2))
            claims.extend([str(i) for i in range(start, end + 1)])
        else:
            # Single claim number
            num_match = re.match(r'(\d+)', part)
            if num_match:
                claims.append(num_match.group(1))
    
    return claims


def _post_process_office_action(result: dict, pdf_text: str, logger) -> dict:
    """
    Post-process Office Action extraction to fix common LLM errors.
    Acts as a safety net — catches issues even when the LLM has the full document.
    
    ENHANCED to extract:
    1. Parent claim assignments (existing)
    2. Examiner phone (existing)
    3. §103 rejection count validation (existing)
    4. Supervisor name and phone (NEW)
    5. Fax number (NEW)
    6. Amended claims (NEW)
    7. IDS submission dates (NEW)
    8. Statutory period and max extension (NEW)
    9. Applicant arguments status (NEW)
    """
    import re

    if not result or not pdf_text:
        return result

    # ══════════════════════════════════════════════════════════════════════════
    # 1. Fix parent claim assignments (existing functionality)
    # ══════════════════════════════════════════════════════════════════════════
    parent_map = {}

    # Pattern: "Claim N:" ... "the [thing] of claim M"
    for match in re.finditer(
        r'Claim\s+(\d+)\s*[:\.].*?the\s+\w[\w\s-]{0,40}?\s+of\s+claim\s+(\d+)',
        pdf_text, re.IGNORECASE | re.DOTALL
    ):
        cn, pn = match.group(1), match.group(2)
        if cn != pn:
            parent_map[cn] = pn

    # Pattern: explicit "Claim X depends on claim Y" (from §112(d) text)
    for match in re.finditer(
        r'[Cc]laim\s+(\d+)\s+depend[s]?\s+(?:on|from)\s+claim\s+(\d+)',
        pdf_text
    ):
        parent_map[match.group(1)] = match.group(2)

    fixes = 0
    for claim in result.get("claims_status", []):
        cn = str(claim.get("claim_number", ""))
        if cn in parent_map and str(claim.get("parent_claim")) != parent_map[cn]:
            old = claim.get("parent_claim")
            claim["parent_claim"] = parent_map[cn]
            claim["dependency_type"] = "Dependent"
            fixes += 1
            logger.info(f"Post-process fix: Claim {cn} parent {old} → {parent_map[cn]}")

    if fixes:
        logger.info(f"Post-processing fixed {fixes} parent claim assignments")

    # ══════════════════════════════════════════════════════════════════════════
    # 2. Extract examiner phone if missing (existing functionality)
    # ══════════════════════════════════════════════════════════════════════════
    header = result.get("header", {})
    if not header.get("examiner_phone"):
        m = re.search(r'\(571\)\s*\d{3}[\s.-]\d{4}', pdf_text)
        if m:
            phone = m.group(0).replace(" ", "")
            header["examiner_phone"] = phone
            logger.info(f"Post-process: extracted examiner phone {phone}")

    # ══════════════════════════════════════════════════════════════════════════
    # 3. Extract SUPERVISOR name and phone (NEW)
    # ══════════════════════════════════════════════════════════════════════════
    if not header.get("supervisor_name"):
        # Pattern 1: Name followed by "Supervisory Patent Examiner"
        m = re.search(
            r'/([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)/?\s*\n?\s*Supervisory\s+Patent\s+Examiner',
            pdf_text, re.IGNORECASE
        )
        if m:
            header["supervisor_name"] = m.group(1).strip()
            logger.info(f"Post-process: extracted supervisor name '{header['supervisor_name']}' (pattern 1)")
        
        # Pattern 2: "[NAME] can be reached on [PHONE]" after "supervisor" mention
        if not header.get("supervisor_name"):
            m = re.search(
                r'(?:supervisor|supervisory)[^\n]{0,100}?([A-Z][A-Z\s]+[A-Z])\s+can\s+be\s+reached\s+on\s+(\d{10})',
                pdf_text, re.IGNORECASE
            )
            if m:
                header["supervisor_name"] = m.group(1).strip().title()
                if not header.get("supervisor_phone"):
                    header["supervisor_phone"] = m.group(2)
                logger.info(f"Post-process: extracted supervisor '{header['supervisor_name']}' phone '{header.get('supervisor_phone')}' (pattern 2)")

        # Pattern 3: Look for signature block with SPE
        if not header.get("supervisor_name"):
            m = re.search(
                r'([A-Z][A-Z\s\.]+[A-Z])\s*,?\s*(?:SPE|Supervisory\s+Patent\s+Examiner)',
                pdf_text
            )
            if m:
                name = m.group(1).strip().title()
                if name.upper() not in ['ART UNIT', 'UNITED STATES', 'PATENT OFFICE']:
                    header["supervisor_name"] = name
                    logger.info(f"Post-process: extracted supervisor name '{name}' (pattern 3)")

    # Extract supervisor phone if we have name but no phone
    if header.get("supervisor_name") and not header.get("supervisor_phone"):
        supervisor_name = header["supervisor_name"].upper()
        m = re.search(
            re.escape(supervisor_name) + r'.*?can\s+be\s+reached\s+on\s+(\d{10})',
            pdf_text.upper(),
            re.DOTALL
        )
        if m:
            header["supervisor_phone"] = m.group(1)
            logger.info(f"Post-process: extracted supervisor phone {header['supervisor_phone']}")

    # ══════════════════════════════════════════════════════════════════════════
    # 4. Extract FAX number (NEW)
    # ══════════════════════════════════════════════════════════════════════════
    if not header.get("fax_number"):
        m = re.search(
            r'fax\s+(?:phone\s+)?number[^\d]{0,30}(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})',
            pdf_text, re.IGNORECASE
        )
        if m:
            digits = re.sub(r'\D', '', m.group(1))
            if len(digits) == 10:
                fax = f"{digits[:3]}-{digits[3:6]}-{digits[6:]}"
                header["fax_number"] = fax
                logger.info(f"Post-process: extracted fax number {fax}")

    # ══════════════════════════════════════════════════════════════════════════
    # 5. Extract AMENDED CLAIMS (NEW)
    # ══════════════════════════════════════════════════════════════════════════
    if not result.get("amended_claims"):
        amended_claims = []
        
        # Pattern 1: "Claims X, Y, Z ... are currently amended"
        m = re.search(
            r'[Cc]laims?\s+([\d,\s\-and]+)\s+(?:is|are)\s+(?:currently\s+)?amended',
            pdf_text
        )
        if m:
            amended_claims.extend(_parse_claim_numbers(m.group(1)))
        
        # Pattern 2: "Currently amended: Claims X-Y"
        m = re.search(
            r'[Cc]urrently\s+[Aa]mended[:\s]+([\d,\s\-and]+)',
            pdf_text
        )
        if m:
            amended_claims.extend(_parse_claim_numbers(m.group(1)))
        
        if amended_claims:
            result["amended_claims"] = sorted(list(set(amended_claims)), key=lambda x: int(x))
            logger.info(f"Post-process: extracted amended claims {result['amended_claims']}")

    # ══════════════════════════════════════════════════════════════════════════
    # 6. Extract IDS SUBMISSION DATES (NEW)
    # ══════════════════════════════════════════════════════════════════════════
    if not result.get("ids_submissions"):
        ids_submissions = []
        
        # Pattern 1: "Information Disclosure Statement(s) ... Paper No(s)/Mail Date X and Y"
        m = re.search(
            r'[Ii]nformation\s+[Dd]isclosure\s+[Ss]tatement.*?'
            r'(?:[Pp]aper\s+[Nn]o[s.]*/)?[Mm]ail\s+[Dd]ate[:\s]*([\d/,\s\-and]+)',
            pdf_text, re.DOTALL
        )
        if m:
            dates = re.findall(r'\d{1,2}/\d{1,2}/\d{4}', m.group(1))
            for date in dates:
                ids_submissions.append({
                    "submission_date": date,
                    "was_considered": True
                })
        
        # Pattern 2: "IDS filed on X has been considered"
        for m in re.finditer(
            r'IDS\s+(?:filed\s+)?(?:on\s+)?(\d{1,2}/\d{1,2}/\d{4})',
            pdf_text, re.IGNORECASE
        ):
            ids_submissions.append({
                "submission_date": m.group(1),
                "was_considered": True
            })
        
        if ids_submissions:
            # Deduplicate by date
            seen_dates = set()
            unique_ids = []
            for ids in ids_submissions:
                if ids["submission_date"] not in seen_dates:
                    seen_dates.add(ids["submission_date"])
                    unique_ids.append(ids)
            result["ids_submissions"] = unique_ids
            logger.info(f"Post-process: extracted IDS submissions {[s['submission_date'] for s in unique_ids]}")

    # ══════════════════════════════════════════════════════════════════════════
    # 7. Extract STATUTORY PERIOD and MAX EXTENSION (NEW)
    # ══════════════════════════════════════════════════════════════════════════
    if not header.get("statutory_period_months"):
        m = re.search(
            r'shortened\s+statutory\s+period.*?(?:expire|reply)[^\d]*'
            r'(?:(\w+)\s+)?\((\d)\)\s*MONTHS?',
            pdf_text, re.IGNORECASE | re.DOTALL
        )
        if m:
            header["statutory_period_months"] = int(m.group(2))
            logger.info(f"Post-process: extracted statutory period {m.group(2)} months")

    if not header.get("max_extension_months"):
        m = re.search(
            r'maximum\s+statutory\s+period.*?(?:expire|after)[^\d]*'
            r'(?:(\w+)\s+)?\((\d)\)\s*MONTHS?',
            pdf_text, re.IGNORECASE | re.DOTALL
        )
        if m:
            header["max_extension_months"] = int(m.group(2))
            logger.info(f"Post-process: extracted max extension {m.group(2)} months")

    # ══════════════════════════════════════════════════════════════════════════
    # 8. Extract APPLICANT ARGUMENTS STATUS (NEW)
    # ══════════════════════════════════════════════════════════════════════════
    if not result.get("applicant_arguments"):
        applicant_arguments = []
        
        # Pattern: "Applicant's arguments regarding claims X, Y, Z are ... moot"
        for m in re.finditer(
            r"[Aa]pplicant'?s?\s+arguments?\s+(?:regarding|with\s+respect\s+to)\s+"
            r"claims?\s+([\d,\s\-and]+)\s+(?:is|are)\s+(?:fully\s+)?(?:considered\s+)?(?:but\s+)?"
            r"(moot|persuasive|not\s+persuasive|unpersuasive)",
            pdf_text, re.IGNORECASE
        ):
            claims = _parse_claim_numbers(m.group(1))
            status = m.group(2).lower().replace(" ", "_")
            if "not" in status or "unpersuasive" in status:
                status = "not_persuasive"
            
            reason = None
            reason_match = re.search(
                r"(moot|not\s+persuasive)\s+(in\s+view\s+of[^.]+)",
                pdf_text[m.start():m.start()+500], re.IGNORECASE
            )
            if reason_match:
                reason = reason_match.group(2).strip()
            
            applicant_arguments.append({
                "status": status,
                "affected_claims": claims,
                "reason": reason
            })
        
        if applicant_arguments:
            result["applicant_arguments"] = applicant_arguments
            logger.info(f"Post-process: extracted {len(applicant_arguments)} applicant argument status entries")

    # ══════════════════════════════════════════════════════════════════════════
    # 9. Sanity-check §103 rejection count (existing functionality)
    # ══════════════════════════════════════════════════════════════════════════
    existing_rejections = result.get("rejections", [])
    existing_103_count = sum(
        1 for r in existing_rejections
        if "103" in str(r.get("rejection_type", ""))
        or "103" in str(r.get("rejection_type_normalized", ""))
    )

    sec103_matches = re.findall(
        r'(?:claims?\s+[\d,\s\-and]+\s+(?:is|are)\s+rejected\s+under\s+'
        r'35\s+U\.?S\.?C\.?\s*(?:§\s*)?103)',
        pdf_text, re.IGNORECASE
    )

    if len(sec103_matches) > existing_103_count:
        logger.warning(
            f"§103 MISMATCH: PDF contains {len(sec103_matches)} §103 rejection blocks "
            f"but LLM only extracted {existing_103_count}. "
            f"Check if page truncation or output token limits are causing this."
        )
        result.setdefault("_extraction_warnings", []).append(
            f"Possible incomplete extraction: {len(sec103_matches)} §103 rejections "
            f"detected in PDF text but only {existing_103_count} were extracted."
        )

    return result

class LLMService:
    def __init__(self):
        self._initialize_client()
        if fitz:
             logger.info("LLMService ready with PyMuPDF support.")
        else:
             logger.warning("LLMService running WITHOUT PyMuPDF. Advanced PDF processing disabled.")

    def _log_token_usage(self, response: Any, operation: str):
        """
        Logs token usage and estimated cost for a Gemini response.
        """
        try:
            if hasattr(response, 'usage_metadata'):
                usage = response.usage_metadata
                prompt_tokens = usage.prompt_token_count
                candidates_tokens = usage.candidates_token_count
                total_tokens = usage.total_token_count
                
                # Calculate estimated cost (based on Gemini 2.5 Pro pricing: ~$1.25/1M input, ~$5.00/1M output)
                input_cost = (prompt_tokens / 1_000_000) * 1.25
                output_cost = (candidates_tokens / 1_000_000) * 5.00
                total_cost = input_cost + output_cost
                
                logger.info(
                    f"Token Usage [{operation}]: Input={prompt_tokens}, Output={candidates_tokens}, Total={total_tokens}",
                    extra={
                        "extra_data": {
                            "token_usage": {
                                "prompt_tokens": prompt_tokens,
                                "completion_tokens": candidates_tokens,
                                "total_tokens": total_tokens,
                                "estimated_cost_usd": round(total_cost, 6)
                            }
                        }
                    }
                )
        except Exception as e:
            logger.warning(f"Failed to log token usage: {e}")

    def _initialize_client(self):
        try:
            logger.info("Attempting to initialize Gemini client...")
            logger.info(f"Using Gemini model: {settings.GEMINI_MODEL}")
            if settings.GOOGLE_API_KEY:
                # Log a masked version of the key to ensure we see it's there
                masked_key = f"{settings.GOOGLE_API_KEY[:4]}...{settings.GOOGLE_API_KEY[-4:]}" if len(settings.GOOGLE_API_KEY) > 8 else "***"
                logger.info(f"GOOGLE_API_KEY found: {masked_key}")
                self.client = genai.Client(api_key=settings.GOOGLE_API_KEY)
                logger.info(f"Initialized Gemini client successfully with model: {settings.GEMINI_MODEL}")
                
                # Test model availability by listing models (optional diagnostic)
                try:
                    logger.info("Testing model availability...")
                    # This is a simple validation that the client works
                    logger.info(f"Client initialized successfully for model: {settings.GEMINI_MODEL}")
                except Exception as test_e:
                    logger.warning(f"Model availability test failed (but client initialized): {test_e}")
                    
            else:
                logger.warning("GOOGLE_API_KEY not found. LLM service not initialized.")
                self.client = None
        except Exception as e:
            logger.error(f"Failed to initialize Gemini client: {e}", exc_info=True)
            self.client = None

    async def upload_file(self, file: Union[str, IO], mime_type: str = "application/pdf"):
        """
        Uploads a file to Gemini for multimodal processing.
        Accepts either a file path (str) or a file-like object (IO).
        """
        if not self.client:
            raise Exception("LLM service not initialized")
        
        try:
            log_name = file if isinstance(file, str) else "memory_stream"
            logger.info(f"Uploading file to Gemini: {log_name}")
            
            # Run in thread pool since library is synchronous
            file_obj = await asyncio.to_thread(
                self.client.files.upload,
                file=file,
                config={'mime_type': mime_type}
            )
            logger.info(f"File uploaded successfully: {file_obj.name}")
            return file_obj
        except Exception as e:
            logger.error(f"Failed to upload file to Gemini: {e}")
            raise e

    async def generate_structured_content(
        self,
        prompt: str,
        file_obj: Any = None,
        schema: Optional[Dict[str, Any]] = None,
        retries: int = 3
    ) -> Dict[str, Any]:
        """
        Generates content from the LLM and parses it as JSON.
        Supports multimodal input (text + file).
        Includes retry logic for transient failures.
        """
        try:
            if not self.client:
                logger.error("LLM Service not initialized when calling generate_structured_content")
                raise Exception("LLM service not initialized")

            # Construct prompt to enforce JSON output
            json_instruction = "\n\nPlease provide the output in valid JSON format."
            if schema:
                json_instruction += f"\nFollow this schema:\n{json.dumps(schema, indent=2)}"
            
            final_text_prompt = prompt + json_instruction

            # Prepare contents
            if file_obj:
                contents = [file_obj, final_text_prompt]
            else:
                contents = final_text_prompt

            for attempt in range(retries):
                try:
                    logger.info(f"Starting LLM generation attempt {attempt + 1}/{retries}")
                    
                    if not final_text_prompt:
                        logger.error("Prompt is empty")
                        raise ValueError("Prompt cannot be empty")

                    # Run sync Gemini call in thread pool
                    start_time = time.time()
                    try:
                        logger.info(f"Calling Gemini API with model: {settings.GEMINI_MODEL}")
                        logger.info(f"API call parameters - Temperature: {settings.GEMINI_TEMPERATURE}, Max tokens: {settings.GEMINI_MAX_OUTPUT_TOKENS}")
                        response = await asyncio.to_thread(
                            self.client.models.generate_content,
                            model=settings.GEMINI_MODEL,
                            contents=contents,
                            config=types.GenerateContentConfig(
                                response_mime_type="application/json",
                                temperature=settings.GEMINI_TEMPERATURE,
                                max_output_tokens=settings.GEMINI_MAX_OUTPUT_TOKENS
                            )
                        )
                        logger.info("Gemini API call returned successfully")
                        
                        # Record latency
                        duration = time.time() - start_time
                        logger.info(f"API call completed in {duration:.2f} seconds")
                        
                        self._log_token_usage(response, "generate_structured_content")
                    except ResourceExhausted as re_err:
                        logger.warning(f"Gemini Rate Limit Exceeded: {re_err}")
                        raise HTTPException(
                            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                            detail="AI Service is currently busy (Rate Limit Exceeded). Please try again in a moment."
                        )
                    except Exception as e:
                        # Enhanced error logging for model-related issues
                        error_msg = str(e)
                        if "NOT_FOUND" in error_msg and "models/" in error_msg:
                            logger.error(f"MODEL ERROR: The model '{settings.GEMINI_MODEL}' was not found or is not supported. Error: {e}")
                            logger.error("Available models may have changed. Consider updating GEMINI_MODEL in configuration.")
                        elif "generateContent" in error_msg:
                            logger.error(f"GENERATE_CONTENT ERROR: The model '{settings.GEMINI_MODEL}' does not support generateContent. Error: {e}")
                        else:
                            logger.error(f"Gemini API execution failed: {e}", exc_info=True)
                        raise e
                    
                    # Log raw response for debugging
                    try:
                        response_text = response.text
                        if not response_text:
                             if hasattr(response, 'candidates') and response.candidates:
                                logger.info(f"Found candidates: {response.candidates}")
                                response_text = response.candidates[0].content.parts[0].text
                             else:
                                raise ValueError("Could not extract text from response")
        
                        # DEBUG: Log the first 500 chars of raw LLM output to see what it generated
                        logger.info(f"RAW LLM RESPONSE (First 500 chars): {response_text[:500]}")
        
                    except Exception as e:
                        logger.error(f"Failed to access response text: {e}", exc_info=True)
                        raise e
        
                    # Parse JSON
                    try:
                        return json.loads(response_text)
                    except json.JSONDecodeError:
                        logger.warning("Initial JSON parse failed, attempting cleanup...")
                        text = response_text
                        
                        # Extract content between code blocks if present
                        if "```" in text:
                            pattern = r"```(?:json)?\s*(.*?)\s*```"
                            match = re.search(pattern, text, re.DOTALL)
                            if match:
                                text = match.group(1)
                        
                        # Find first { and last }
                        start = text.find('{')
                        end = text.rfind('}')
                        if start != -1 and end != -1:
                            text = text[start:end+1]
                            
                        return json.loads(text)
                        
                except Exception as e:
                    logger.warning(f"LLM generation failed (attempt {attempt + 1}/{retries}): {e}")
                    if attempt == retries - 1:
                        raise e
                    wait_time = (2 ** attempt) * 2  # Exponential backoff
                    await asyncio.sleep(wait_time)
        except Exception as outer_e:
            logger.critical(f"CRITICAL ERROR in generate_structured_content: {outer_e}", exc_info=True)
            raise outer_e

    async def analyze_cover_sheet(
        self,
        file_path: str,
        file_content: Optional[bytes] = None,
        extraction_strategy: str = "text",  # NEW: "text", "vision", or "docx_text"
        progress_callback: Optional[Callable[[int, str], Awaitable[None]]] = None
    ) -> PatentApplicationMetadata:
        """
        Analyzes the cover sheet PDF with Parallel Execution optimization.
        Run Local XFA Check AND Remote File Upload simultaneously to minimize latency.
        
        Args:
            file_path: Path to the file (used for logging/extension even if content is provided)
            file_content: Optional raw bytes of the file. If provided, avoids disk reads.
            extraction_strategy: "text", "vision", or "docx_text" - routing hint from validation
            progress_callback: Optional callback for status updates
        """
        file_ext = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else 'pdf'

        # ── Route based on strategy ──────────────────────────────────
        if file_ext == 'docx' or extraction_strategy == "docx_text":
            return await self._analyze_docx_document(file_path, file_content, progress_callback)

        # P1: If pre-check determined this is a scanned PDF, skip text extraction
        if extraction_strategy == "vision":
            logger.info("Scanned PDF detected — using vision extraction directly")
            return await self._analyze_pdf_direct_fallback(
                file_path, file_content=file_content, progress_callback=progress_callback
            )
        # ─────────────────────────────────────────────────────────────

        # ── Continue with existing PDF logic below ──
        logger.info(f"--- ANALYZING PDF WITH GEMINI: {file_path} ---")
        logger.info(f"Concurrency Limit: {settings.MAX_CONCURRENT_EXTRACTIONS}")
        
        # Prepare upload source (BytesIO or Path)
        if file_content:
            upload_source = io.BytesIO(file_content)
        else:
            upload_source = file_path

        if progress_callback:
            logger.info("Reporting progress: 10%")
            await progress_callback(10, "Initiating parallel analysis...")

        # Determine page count to decide strategy
        page_count = 0
        try:
            if file_content:
                reader = PdfReader(io.BytesIO(file_content))
            else:
                reader = PdfReader(file_path)

            # ── P0: Handle encryption at LLM layer too (belt + suspenders) ──
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    from app.services.file_validators import FileValidationError
                    raise FileValidationError(
                        "PDF is encrypted and cannot be processed.",
                        error_code="PDF_ENCRYPTED",
                    )
            # ────────────────────────────────────────────────────────────────

            page_count = len(reader.pages)
            logger.info(f"PDF Page Count: {page_count}")
        except Exception as e:
            logger.warning(f"Failed to get page count: {e}")

        # Standard path: try text extraction first
        try:
            text_start = datetime.utcnow()
            text_content = await self._extract_text_locally(file_path, file_content)
            
            # ══════════════════════════════════════════════════════════════════
            # APPLY FRAGMENTED TEXT CLEANING - Fix for word-per-line PDFs
            # ══════════════════════════════════════════════════════════════════
            original_length = len(text_content)
            text_content = clean_fragmented_text(text_content)
            cleaned_length = len(text_content)
            
            if original_length != cleaned_length:
                logger.info(
                    f"Text cleaned: {original_length} -> {cleaned_length} chars "
                    f"(removed {original_length - cleaned_length} fragmentation artifacts)"
                )
            # ══════════════════════════════════════════════════════════════════
            
            # Check if text is sufficient (not just empty pages or headers)
            # We look for a reasonable amount of text or specific form markers
            # Remove standard markers to see if there's actual content
            clean_text = re.sub(r'--- PAGE \d+ ---', '', text_content)
            clean_text = clean_text.replace("--- FORM FIELD DATA", "").replace("--- END FORM DATA ---", "")
            clean_text = clean_text.replace("[EMPTY PAGE TEXT - LIKELY IMAGE OR XFA]", "")
            clean_text = clean_text.strip()
            
            if len(clean_text) > 100:
                logger.info(f"Text-First Strategy: Sufficient text found ({len(clean_text)} chars). Skipping upload.")
                
                if progress_callback:
                    await progress_callback(30, "Analyzing extracted text...")
                    
                result = await self._analyze_text_only(text_content)
                
                # Basic validation: ensure we got something
                if result.title or result.application_number or (result.inventors and len(result.inventors) > 0):
                     logger.info(f"Text-First Analysis Successful. Latency: {(datetime.utcnow() - text_start).total_seconds()}s")
                     return result
                else:
                    logger.warning("Text-First Analysis returned empty data. Falling back to Vision.")
            else:
                logger.info("Insufficient text extracted, falling back to vision")
                return await self._analyze_pdf_direct_fallback(
                    file_path, file_content=file_content, progress_callback=progress_callback
                )

        except Exception as e:
            # ── P0: Catch any PdfReader crash gracefully ─────────────
            logger.error(f"PDF reading failed: {e}")
            # Try vision fallback before giving up
            try:
                return await self._analyze_pdf_direct_fallback(
                    file_path, file_content=file_content, progress_callback=progress_callback
                )
            except Exception as fallback_error:
                logger.error(f"Vision fallback also failed: {fallback_error}")
                raise Exception(
                    "Unable to extract information from this PDF. "
                    "The file may be corrupted or in an unsupported format."
                )
            # ─────────────────────────────────────────────────────────

        # FALLBACK: Vision / Native PDF (requires upload)
        logger.info("Initiating file upload for Vision analysis...")
        if progress_callback:
             await progress_callback(40, "Uploading document for Vision analysis...")

        upload_task = asyncio.create_task(self.upload_file(upload_source))
        
        # STRATEGY 2: Check for XFA Dynamic Form Data (Local CPU) - While uploading
        try:
            xfa_start = datetime.utcnow()
            xfa_data = await self._extract_xfa_data(file_path, file_content)
            logger.info(f"XFA Check took: {(datetime.utcnow() - xfa_start).total_seconds()}s")
            
            if xfa_data:
                logger.info("XFA Dynamic Form detected! Using direct XML extraction path.")
                xfa_result = await self._analyze_xfa_xml(xfa_data)
                
                # Validation
                if xfa_result.inventors and len(xfa_result.inventors) > 0:
                     valid_inventors = [i for i in xfa_result.inventors if i.name or i.last_name]
                     if valid_inventors:
                         logger.info(f"Successfully extracted {len(valid_inventors)} inventors from XFA data.")
                         # Cancel upload as it's not needed
                         upload_task.cancel()
                         try:
                             await upload_task
                         except asyncio.CancelledError:
                             pass
                         xfa_result.inventors = valid_inventors
                         return xfa_result
        except Exception as e:
            logger.warning(f"XFA detection failed (continuing to vision fallback): {e}")

        # STRATEGY 2: Fast-Track (Native PDF) using pre-started upload
        # Use ONLY for small documents (< 50 pages)
        if page_count < 50:
            logger.info("Document is small (< 50 pages). Using Native PDF Fast-Track strategy...")
            
            if progress_callback:
                await progress_callback(20, "Analyzing full document (Fast-Track)...")
                
            try:
                # Wait for upload to complete (if not already)
                upload_start = datetime.utcnow()
                file_obj = await upload_task
                logger.info(f"File upload ready. Total upload wait: {(datetime.utcnow() - upload_start).total_seconds()}s")
                
                return await self._analyze_pdf_direct_fallback(file_path, file_obj=file_obj, file_content=file_content)
            except Exception as e:
                logger.warning(f"Native PDF extraction failed: {e}. Falling back to Hybrid Page-by-Page strategy.")
        else:
            logger.info(f"Document is large ({page_count} pages). Skipping Native PDF to use Hybrid Parallel strategy.")

        # STRATEGY 3: Unified Chunking Strategy (Vision) - Fallback
        # This replaces the old "Page-by-Page Image" strategy with a robust "PDF Chunking" approach.
        logger.info("Falling back to Unified Chunking Strategy (Vision)...")
        
        if progress_callback:
            logger.info("Reporting progress: 20%")
            await progress_callback(20, "Analyzing document chunks with Vision...")

        try:
            # We need raw bytes for chunking
            if not file_content:
                with open(file_path, "rb") as f:
                    pdf_bytes = f.read()
            else:
                pdf_bytes = file_content
            
            chunk_result = await self._analyze_document_chunked_structured(
                file_bytes=pdf_bytes,
                filename=os.path.basename(file_path),
                total_pages=page_count,
                progress_callback=progress_callback
            )

            # STRATEGY 4: Final Fallback - Direct PDF Upload
            # If chunking found literally nothing (or failed), try one last desperate Direct Upload
            # But only if the chunking result is basically empty
            if not chunk_result.inventors and not chunk_result.title:
                logger.warning("⚠️ Chunking found no metadata. Attempting Final Fallback: Direct PDF Upload...")
                return await self._analyze_pdf_direct_fallback(file_path, file_content=file_content)
            
            return chunk_result

        except Exception as e:
            logger.error(f"Unified Chunking analysis failed: {e}")
            # Final fallback if chunking crashes completely
            return await self._analyze_pdf_direct_fallback(file_path, file_content=file_content)

    async def _extract_xfa_data(self, file_path: str, file_content: Optional[bytes] = None) -> Optional[str]:
        """
        Checks if the PDF is an XFA form and extracts the internal XML data.
        """
        def _read_xfa():
            try:
                if file_content:
                    reader = PdfReader(io.BytesIO(file_content))
                else:
                    reader = PdfReader(file_path)
                    
                if "/AcroForm" in reader.trailer["/Root"]:
                    acroform = reader.trailer["/Root"]["/AcroForm"]
                    if "/XFA" in acroform:
                        # XFA content can be a list or a stream
                        xfa = acroform["/XFA"]
                        # Often it's a list of [key, indirect_object, key, indirect_object...]
                        # We want to find the 'datasets' packet usually
                        
                        # Targeted Extraction: Prioritize 'datasets' which contains user data
                        xml_content = []
                        
                        if isinstance(xfa, list):
                            # XFA is a list of keys and values: [key1, val1, key2, val2...]
                            # We want to grab everything, but prioritize 'datasets'
                            for i in range(0, len(xfa), 2):
                                key = xfa[i]
                                obj = xfa[i+1]
                                
                                # We specifically want the 'datasets' packet as it contains the actual USER DATA.
                                # IMPORTANT: 'template' contains the empty form structure (400KB+) which confuses the LLM.
                                # We ONLY want 'datasets' to give the LLM pure data.
                                if key == 'datasets':
                                    try:
                                        data = obj.get_object().get_data()
                                        if data:
                                            decoded_data = data.decode('utf-8', errors='ignore')
                                            xml_content.append(f"<!-- {key} START -->")
                                            xml_content.append(decoded_data)
                                            xml_content.append(f"<!-- {key} END -->")
                                    except Exception as e:
                                        logger.warning(f"Failed to read XFA packet {key}: {e}")
                                        
                        else:
                            # Single stream fallback
                            try:
                                xml_content.append(xfa.get_object().get_data().decode('utf-8', errors='ignore'))
                            except:
                                pass
                        
                        full_xml = "\n".join(xml_content)
                        if len(full_xml) > 100:
                            logger.info(f"Successfully extracted XFA XML (Length: {len(full_xml)} bytes)")
                            return full_xml
                return None
            except Exception as e:
                logger.warning(f"Error reading XFA data: {e}")
                return None

        return await asyncio.to_thread(_read_xfa)

    async def _analyze_form_text(self, form_text: str) -> PatentApplicationMetadata:
        """
        Analyzes raw form field text extracted by pypdf.
        """
        prompt = f"""
        Analyze the provided PDF Form Data (Key-Value pairs) from a Patent Application.
        Extract the patent metadata by inferring the meaning of the field keys and values.
        
        ## FORM DATA
        {form_text[:50000]}
        
        ## INSTRUCTIONS
        - The data is presented as 'Field_Name: Value'.
        - Look for keys like 'Title', 'InventionTitle', 'ApplicationNo', 'AppNum', etc.
        - **Inventors**: Look for repeating fields like 'GivenName_1', 'FamilyName_1', 'Address_1' etc.
        - Reconstruct the inventor objects from these flattened keys.
        - **Applicant/Company**: Look for fields like 'ApplicantName', 'CompanyName', 'OrgName', 'AssigneeName', etc.
        - Extract applicant address fields like 'ApplicantAddress', 'CompanyAddress', 'OrgAddress', etc.
        
        ## OUTPUT SCHEMA
        Return JSON with:
        - _debug_reasoning (string)
        - title
        - application_number
        - entity_status
        - inventors (list of objects)
        - applicant (object with company/organization information)
        """
        
        schema = {
            "_debug_reasoning": "Explain which keys were mapped to which fields",
            "title": "Title found (or null)",
            "application_number": "Application number (or null)",
            "entity_status": "Entity status (or null)",
            "inventors": [
                {
                    "name": "Full Name",
                    "first_name": "First name",
                    "middle_name": "Middle name",
                    "last_name": "Last name",
                    "city": "City",
                    "state": "State",
                    "country": "Country",
                    "street_address": "Street address",
                    "full_address": "Full address string"
                }
            ],
            "applicant": {
                "name": "Company/Applicant name",
                "street_address": "Street address",
                "city": "City",
                "state": "State",
                "zip_code": "Postal/ZIP code",
                "country": "Country"
            }
        }
        
        result = await self.generate_structured_content(prompt=prompt, schema=schema)
        
        # Post-processing
        if result.get("inventors"):
            for inventor in result["inventors"]:
                if inventor.get("name") and not inventor.get("last_name"):
                    parts = inventor["name"].split()
                    if len(parts) >= 2:
                        inventor["first_name"] = parts[0]
                        inventor["last_name"] = parts[-1]
        
        return PatentApplicationMetadata(**result)

    async def _analyze_xfa_xml(self, xfa_xml: str) -> PatentApplicationMetadata:
        """
        Analyzes the raw XFA XML data to extract metadata.
        """
        # Truncate XML if it's massive to avoid context limits (though rare for ADS)
        truncated_xml = xfa_xml[:50000]
        
        prompt = f"""
        Analyze the provided XFA Form XML Data from a Patent Application Data Sheet (ADS).
        Extract the patent metadata directly from the XML structure.
        
        ## XML DATA
        {truncated_xml}
        
        ## INSTRUCTIONS
        - The data is structured in XML tags. Look for:
          - Title of Invention
          - Application Number / Control Number
          - Inventor Information (Names, Cities, States, Addresses)
          - Applicant/Company Information (Organization name, address)
        - **Inventors**: Extract ALL inventors found in the XML datasets.
        - **Applicant**: Look for organization/company information in assignee or applicant sections.
        
        ## OUTPUT SCHEMA
        Return JSON with:
        - _debug_reasoning (string)
        - title
        - application_number
        - entity_status
        - inventors (list of objects)
        - applicant (object with company/organization information)
        """
        
        schema = {
            "_debug_reasoning": "Explain where in the XML the data was found",
            "title": "Title found (or null)",
            "application_number": "Application number (or null)",
            "entity_status": "Entity status (or null)",
            "inventors": [
                {
                    "name": "Full Name",
                    "first_name": "First name",
                    "middle_name": "Middle name",
                    "last_name": "Last name",
                    "city": "City",
                    "state": "State",
                    "country": "Country",
                    "street_address": "Street address",
                    "full_address": "Full address string"
                }
            ],
            "applicant": {
                "name": "Company/Applicant name",
                "street_address": "Street address",
                "city": "City",
                "state": "State",
                "zip_code": "Postal/ZIP code",
                "country": "Country"
            }
        }
        
        result = await self.generate_structured_content(prompt=prompt, schema=schema)
        
        # Post-processing same as before
        if result.get("inventors"):
            for inventor in result["inventors"]:
                if inventor.get("name") and not inventor.get("last_name"):
                    parts = inventor["name"].split()
                    if len(parts) >= 2:
                        inventor["first_name"] = parts[0]
                        inventor["last_name"] = parts[-1]
        
        return PatentApplicationMetadata(**result)

    async def _analyze_text_only(self, text_content: str) -> PatentApplicationMetadata:
        """
        Analyzes raw text content to extract metadata.
        Used for Text-First strategy (PDFs with extractable text) AND for DOCX files.
        
        UPDATED: Now also extracts correspondence address, application type,
        and suggested representative figure.
        """
        # ══════════════════════════════════════════════════════════════════
        # APPLY FINAL TEXT CLEANING - Ensure clean text reaches the LLM
        # ══════════════════════════════════════════════════════════════════
        original_length = len(text_content)
        text_content = clean_fragmented_text(text_content)
        cleaned_length = len(text_content)
        
        if original_length != cleaned_length:
            logger.info(
                f"Final text cleaning in _analyze_text_only: {original_length} -> {cleaned_length} chars "
                f"(removed {original_length - cleaned_length} fragmentation artifacts)"
            )
        # ══════════════════════════════════════════════════════════════════
        
        prompt = f"""
        Analyze the provided Text Content from a Patent Application Data Sheet (ADS), cover sheet,
        or patent application document (which may be a DOCX file).
        Extract the patent metadata directly from the text.
        
        ## TEXT CONTENT
        {text_content[:80000]}
        
        ## ╔══════════════════════════════════════════════════════════════════════════════╗
        ## ║  CRITICAL: INVENTOR EXTRACTION RULES - READ CAREFULLY                        ║
        ## ╚══════════════════════════════════════════════════════════════════════════════╝
        
        **ABSOLUTE REQUIREMENT**: You MUST extract EVERY SINGLE inventor listed in the document.
        
        1. **START FROM THE FIRST INVENTOR**: The very first name listed after "Inventors:" or similar
           header MUST be included. Do NOT skip the first inventor.
        
        2. **COUNT THE INVENTORS**: Before extracting, count how many inventors are listed.
           Then verify your output contains that exact number.
        
        3. **HANDLE NAME SUFFIXES**: Names may include suffixes like:
           - "Jr." or "Jr" (Junior)
           - "Sr." or "Sr" (Senior)
           - "III", "IV" (Roman numerals)
           - "Ph.D.", "M.D.", "Esq." (Professional titles)
           These are PART OF THE NAME and should be extracted in the "suffix" field.
        
        4. **EXAMPLE**: If the document lists:
           "Inventors:
            Robert James Smith Jr.
            María Elena García-López
            Wei-Lin Chen"
           
           You MUST return 3 inventors, starting with Robert James Smith Jr.
        
        5. **VERIFICATION**: After extraction, re-read the inventor list and confirm:
           - First inventor matches the first name in the document
           - Total count matches the document
           - No inventors were skipped
        
        ## STANDARD INSTRUCTIONS
        - **Title**: Look for "Title of Invention" or similar headers.
        - **Application Number**: Look for "Application Number", "Control Number".
        - **Entity Status**: Look for indicators like "Small Entity", "Micro Entity", "Large Entity".
        - **Total Drawing Sheets**: ONLY extract if the document explicitly states a specific
          number (e.g., "Total Number of Drawing Sheets: 14" or a filled form field with a
          number). Do NOT count or estimate from figure descriptions, figure labels, or
          drawing references in the text. If not explicitly stated, return null.
        
        ## INVENTOR FIELD DETAILS
        - **first_name**: Given name (e.g., "Robert")
        - **middle_name**: Middle name(s) (e.g., "James") - DO NOT TRUNCATE
        - **last_name**: Family/surname (e.g., "Smith")
        - **suffix**: Name suffix (e.g., "Jr.", "III", "Ph.D.") - EXTRACT THIS SEPARATELY
        - **city**: City of residence
        - **state**: State/province
        - **country**: Country
        - **citizenship**: Country of citizenship (REQUIRED if available)
        - **street_address**: Full mailing address
        - **zip_code**: Complete postal code (REQUIRED if available)
        
        ## MULTI-APPLICANT INSTRUCTIONS
        - **FIND ALL APPLICANTS**: Look for multiple applicants/companies - there may be 2, 3, or more
        - **SEARCH THOROUGHLY**: Check entire document for "Applicant 1:", "Applicant 2:", etc.
        - **COMPANY INDICATORS**: Look for legal suffixes like "LLC", "Inc.", "Corp.", "Ltd.", "GmbH"
        - **DO NOT STOP**: After finding one applicant, actively search for more
        
        ## ═══ NEW FIELDS — CRITICAL EXTRACTION ═══
        
        7. **Correspondence Address** (CRITICAL):
           - Look for "Correspondence Address", "Attorney Address", "Send Correspondence To",
             "Address for Correspondence", law firm information, or patent agent details.
           - This is typically a LAW FIRM or patent attorney address used for USPTO communications.
           - Extract ALL of: firm/person name, street address, city, state, postal code, country,
             phone number, fax, email.
           - Look for patterns: "LLP", "P.C.", "& Associates", law firm names.
           - May also include a USPTO "Customer Number" (5-6 digit number).
           - If correspondence appears alongside attorney info, extract the address.
        
        8. **Application Type** (CRITICAL):
           - Determine the patent application type from the document content.
           - Possible values: "utility", "design", "plant", "provisional", "reissue"
           - **Utility** (most common): System/method/process/apparatus claims, technical specifications.
           - **Design**: Ornamental design claims, design drawings, "D" prefix on app number.
           - **Plant**: Plant variety claims, botanical descriptions.
           - **Provisional**: Explicitly marked "Provisional Application" or "60/" series number.
           - **Reissue**: Marked "Reissue Application" or references an existing patent to be reissued.
           - If not explicitly stated, infer from the claim structure and document content.
           - Most patent applications with system/method claims are "utility".
        
        9. **Suggested Representative Figure** (CRITICAL):
           - Look for "Representative Drawing", "Suggested Drawing Figure", "Representative Figure".
           - If not explicitly stated, analyze the figures/drawings described:
             - The main system overview or architecture diagram is typically Figure 1.
             - Suggest the figure that best represents the overall invention.
           - Extract JUST the figure number: "1", "2", "3A", etc.
           - If no figures are described, return null.
        
        ## OUTPUT SCHEMA
        Return JSON with:
        - _debug_reasoning (string explaining extraction logic, how many applicants found,
          and what correspondence/type/figure info was found)
        - title
        - application_number
        - entity_status
        - total_drawing_sheets
        - inventors (list of objects with postal codes)
        - applicants (list of ALL applicant objects found)
        - correspondence_address (object with firm name, address, phone, etc.)
        - application_type (string: utility/design/plant/provisional/reissue)
        - suggested_figure (string: figure number like "1", "2A", or null)
        """
        
        schema = {
            "_debug_reasoning": "REQUIRED: State 'Found X inventors, first inventor is [name], last inventor is [name]'",
            "title": "Title found (or null)",
            "application_number": "Application number (or null)",
            "entity_status": "Entity status (or null)",
            "total_drawing_sheets": "ONLY if explicitly stated as a number in the document. Do NOT count figures or estimate. Return null if not found.",
            "inventors": [
                {
                    "name": "Full Name (for reference)",
                    "first_name": "First/given name (REQUIRED)",
                    "middle_name": "Complete middle name - DO NOT TRUNCATE",
                    "last_name": "Last/family name (REQUIRED)",
                    "suffix": "Name suffix: Jr., Sr., III, Ph.D., etc. (or null)",
                    "city": "City",
                    "state": "State",
                    "country": "Country",
                    "citizenship": "Citizenship (REQUIRED if available)",
                    "street_address": "Street address",
                    "zip_code": "Complete postal code (REQUIRED if available)",
                    "full_address": "Full address string"
                }
            ],
            "applicants": [
                {
                    "name": "Company/Applicant name",
                    "street_address": "Street address",
                    "city": "City",
                    "state": "State",
                    "zip_code": "Postal/ZIP code",
                    "country": "Country"
                }
            ],
            "correspondence_address": {
                "name": "Law firm or person name (e.g. 'Blakely, Sokoloff, Taylor & Zafman LLP')",
                "address1": "Street address line 1",
                "address2": "Street address line 2 (or null)",
                "city": "City",
                "state": "State (2-letter code for US)",
                "country": "Country",
                "postcode": "Postal/ZIP code",
                "phone": "Phone number including area code",
                "fax": "Fax number (or null)",
                "email": "Email address (or null)",
                "customer_number": "USPTO customer number if found (or null)"
            },
            "application_type": "One of: utility, design, plant, provisional, reissue (or null if cannot determine)",
            "suggested_figure": "Representative figure number as string (e.g. '1', '2A') or null"
        }
        
        result = await self.generate_structured_content(prompt=prompt, schema=schema)
        
        # ══════════════════════════════════════════════════════════════════════════
        # NEW: Log inventor extraction details for debugging
        # ══════════════════════════════════════════════════════════════════════════
        if result.get("_debug_reasoning"):
            logger.info(f"LLM Debug Reasoning: {result.get('_debug_reasoning')}")

        if result.get("inventors"):
            logger.info(f"Extracted {len(result['inventors'])} inventors:")
            for i, inventor in enumerate(result["inventors"]):
                full_name = f"{inventor.get('first_name', '')} {inventor.get('middle_name', '')} {inventor.get('last_name', '')} {inventor.get('suffix', '')}".strip()
                logger.info(f"  Inventor {i+1}: {full_name}")
        
        # Post-processing for name splitting (same as before)
        if result.get("inventors"):
            for inventor in result["inventors"]:
                if inventor.get("name") and not inventor.get("last_name"):
                    parts = inventor["name"].split()
                    if len(parts) >= 2:
                        inventor["first_name"] = parts[0]
                        inventor["last_name"] = parts[-1]

        # Post-processing: normalize application_type to lowercase
        if result.get("application_type"):
            result["application_type"] = result["application_type"].strip().lower()

        # Post-processing: normalize suggested_figure to string
        if result.get("suggested_figure") is not None:
            result["suggested_figure"] = str(result["suggested_figure"]).strip()

        return PatentApplicationMetadata(**result)

    async def _analyze_single_page_image(self, img_path: str, page_num: int, page_text: str = "") -> Dict[str, Any]:
        """
        Analyzes a single page image AND its text content to extract partial metadata.
        """
        try:
            file_obj = await self.upload_file(img_path, mime_type="image/jpeg")
            
            # ══════════════════════════════════════════════════════════════════
            # APPLY TEXT CLEANING to page text before sending to LLM
            # ══════════════════════════════════════════════════════════════════
            if page_text:
                original_page_text = page_text
                page_text = clean_fragmented_text(page_text)
                if len(original_page_text) != len(page_text):
                    logger.debug(f"Page {page_num} text cleaned for image analysis: {len(original_page_text)} -> {len(page_text)} chars")
            # ══════════════════════════════════════════════════════════════════
            
            prompt = f"""
            Analyze this specific page (Page {page_num}) of a Patent Application Data Sheet (ADS).
            I am providing BOTH the visual image AND the raw text content for this page.
            
            ## RAW TEXT CONTENT
            {page_text[:10000]} # Limit text to avoid context overflow if huge
            
            ## INSTRUCTIONS
            1. **Visual Reasoning**: First, explain what you see on the page in the '_debug_reasoning' field.
               - Do you see an "Inventor Information" header?
               - Do you see a table structure?
               - Does the raw text contain names that might be illegible in the image?
            2. **Inventors Extraction**:
               - **COMBINE SOURCES**: Use the Image to understand the layout (rows/columns) and the Text to get accurate spelling.
               - **SEARCH AGGRESSIVELY**: Look for *any* blocks that contain names and addresses.
               - **Address**: If you can't separate City/State, just put the whole address in 'full_address' or 'street_address'.
            3. **Header Info**: Look for Title, Application Number, Entity Status.

            ## OUTPUT SCHEMA
            Return JSON with:
            - _debug_reasoning (string): Description of page content and logic used.
            - title (string/null)
            - application_number (string/null)
            - entity_status (string/null)
            - inventors (list of objects)
            """
            
            schema = {
                "_debug_reasoning": "Explain what sections were found on this page (e.g., 'Found Inventor Info table with 2 rows')",
                "title": "Title found on this page (or null)",
                "application_number": "Application number found on this page (or null)",
                "entity_status": "Entity status found on this page (or null)",
                "inventors": [
                    {
                        "name": "Full Name",
                        "first_name": "First name",
                        "middle_name": "Middle name",
                        "last_name": "Last name",
                        "city": "City",
                        "state": "State",
                        "country": "Country",
                        "street_address": "Street address / Mailing address",
                        "full_address": "Full address string (fallback)"
                    }
                ]
            }

            result = await self.generate_structured_content(
                prompt=prompt,
                file_obj=file_obj,
                schema=schema
            )
            
            # Log the reasoning for debugging purposes
            if result.get("_debug_reasoning"):
                logger.info(f"Page {page_num} Analysis: {result.get('_debug_reasoning')}")
            
            return result
            
        except Exception as e:
            logger.warning(f"Failed to analyze page {page_num}: {e}")
            return {}

    async def _analyze_pdf_direct_fallback(self, file_path: str, file_obj: Any = None, file_content: Optional[bytes] = None, progress_callback: Optional[Callable[[int, str], Awaitable[None]]] = None) -> PatentApplicationMetadata:
        """
        Single-pass native PDF extraction.
        Accepts optional pre-uploaded file_obj to save time.
        """
        # Upload file to Gemini if not provided
        if not file_obj:
            try:
                upload_source = io.BytesIO(file_content) if file_content else file_path
                file_obj = await self.upload_file(upload_source)
            except Exception as e:
                logger.error(f"Failed to upload file for analysis: {e}")
                raise e

        # Construct prompt for direct visual extraction
        prompt = """
        Analyze the provided document, which is likely a **Patent Application Data Sheet (ADS)** or similar cover sheet.
        Your goal is to extract specific bibliographic data with HIGH PRECISION.

        ## ╔══════════════════════════════════════════════════════════════════════════════╗
        ## ║  CRITICAL: INVENTOR EXTRACTION RULES - READ THIS FIRST                       ║
        ## ╚══════════════════════════════════════════════════════════════════════════════╝
        
        **ABSOLUTE REQUIREMENT**: Extract EVERY SINGLE inventor. Do NOT skip any.
        
        1. **START FROM THE FIRST INVENTOR**: The very first name in the inventor list MUST be included.
        
        2. **COUNT BEFORE EXTRACTING**:
           - First, scan the entire document and COUNT how many inventors are listed
           - Then extract each one, starting from the first
           - Verify your output count matches
        
        3. **HANDLE NAME SUFFIXES**: Names may include:
           - "Jr." / "Jr" / "Junior"
           - "Sr." / "Sr" / "Senior"
           - "III", "IV", "V" (Roman numerals)
           - "Ph.D.", "M.D.", "Esq."
           Extract these in the "suffix" field, NOT as part of last_name.
        
        4. **EXAMPLE**: If document shows:
           "1. Robert James Smith Jr.    2. Maria Garcia    3. John Chen"
           You MUST return exactly 3 inventors, with Robert James Smith Jr. as the FIRST one.
        
        5. **MULTI-PAGE CHECK**: Inventor lists often SPAN MULTIPLE PAGES. Check ALL pages.

        ## DOCUMENT STRUCTURE AWARENESS
        - **ADS Forms (PTO/AIA/14)**: Use structured tables
        - **Inventor Section**: Look for "Inventor Information" header
        - **Columns**: Names may be split into "Given Name", "Middle Name", "Family Name"

        ## EXTRACTION INSTRUCTIONS
        1. **Title**: Extract the "Title of Invention"
        2. **Application Number**: Extract if present
        3. **Filing Date**: Extract if present
        4. **Entity Status**: "Small Entity", "Micro Entity", or "Large Entity"
        5. **Total Drawing Sheets**: ONLY if explicitly stated as a number
        6. **Inventors**: Extract ALL with complete information including suffix
        7. **Applicant/Company**: Look for company names, assignee information
        8. **Correspondence Address**: Law firm/attorney contact info
        9. **Application Type**: utility, design, plant, provisional, or reissue
        10. **Suggested Representative Figure**: Figure number

        ## DATA CLEANING RULES
        - Remove legal boilerplate
        - If a field is empty, return null
        - Do NOT Hallucinate - only extract what is visible

        The output must be valid JSON matching the provided schema.
        """
        
        schema = {
            "_debug_reasoning": "MUST state: 'Found [N] inventors. First inventor: [name]. Last inventor: [name].'",
            "title": "Title of the invention",
            "application_number": "Application number",
            "filing_date": "Filing date",
            "entity_status": "Entity status",
            "total_drawing_sheets": "Only if explicitly stated. Return null if not found.",
            "inventors": [
                {
                    "name": "Full Name (reference only)",
                    "first_name": "First name (REQUIRED)",
                    "middle_name": "Complete middle name - DO NOT TRUNCATE",
                    "last_name": "Last name (REQUIRED)",
                    "suffix": "Jr., Sr., III, Ph.D., etc. (or null)",
                    "city": "City",
                    "state": "State",
                    "country": "Country",
                    "citizenship": "Citizenship (REQUIRED)",
                    "street_address": "Street address"
                }
            ],
            "applicants": [
                {
                    "name": "Company/Applicant name",
                    "street_address": "Street address",
                    "city": "City",
                    "state": "State",
                    "zip_code": "Postal/ZIP code",
                    "country": "Country"
                }
            ],
            "correspondence_address": {
                "name": "Law firm or person name",
                "address1": "Street address line 1",
                "address2": "Street address line 2 (or null)",
                "city": "City",
                "state": "State",
                "country": "Country",
                "postcode": "Postal/ZIP code",
                "phone": "Phone number",
                "fax": "Fax number (or null)",
                "email": "Email address (or null)",
                "customer_number": "USPTO customer number (or null)"
            },
            "application_type": "utility, design, plant, provisional, or reissue",
            "suggested_figure": "Representative figure number (e.g. '1') or null"
        }
        
        try:
            # Pass the file object DIRECTLY to the LLM along with the prompt
            result = await self.generate_structured_content(
                prompt=prompt,
                file_obj=file_obj,  # <--- Key change: Passing the file object
                schema=schema
            )
            
            # Validate that we actually got meaningful data
            if not result:
                raise ValueError("LLM returned empty response")
            
            # ══════════════════════════════════════════════════════════════════════════
            # POST-PROCESSING: Log and verify inventor extraction
            # ══════════════════════════════════════════════════════════════════════════
            if result.get("_debug_reasoning"):
                logger.info(f"LLM Debug (Vision): {result.get('_debug_reasoning')}")
            
            if result.get("inventors"):
                logger.info(f"Vision extraction found {len(result['inventors'])} inventors:")
                for i, inventor in enumerate(result["inventors"]):
                    full_name = f"{inventor.get('first_name', '')} {inventor.get('middle_name', '')} {inventor.get('last_name', '')} {inventor.get('suffix', '')}".strip()
                    logger.info(f"  Inventor {i+1}: {full_name}")
                    
                    # Split name if only 'name' provided
                    if inventor.get("name") and not inventor.get("last_name"):
                        parts = inventor["name"].split()
                        if len(parts) >= 2:
                            inventor["first_name"] = parts[0]
                            inventor["last_name"] = parts[-1]
                            if len(parts) > 2:
                                inventor["middle_name"] = " ".join(parts[1:-1])
                        elif len(parts) == 1:
                            inventor["first_name"] = parts[0]

            # Normalize application_type
            if result.get("application_type"):
                result["application_type"] = result["application_type"].strip().lower()

            # Normalize suggested_figure to string
            if result.get("suggested_figure") is not None:
                result["suggested_figure"] = str(result["suggested_figure"]).strip()

            return PatentApplicationMetadata(**result)
            
        except Exception as e:
            logger.error(f"Error analyzing cover sheet: {e}")
            raise e

    async def _convert_pdf_to_images(self, file_path: str, file_content: Optional[bytes] = None) -> List[str]:
        """
        Converts PDF pages to JPEG images using PyMuPDF (fitz).
        Returns a list of temporary file paths.
        """
        if fitz is None:
            logger.error("PyMuPDF (fitz) is not installed. Image conversion fallback unavailable.")
            return []

        def _convert():
            image_paths = []
            try:
                if file_content:
                    doc = fitz.open(stream=file_content, filetype="pdf")
                    # If we don't have a real path, create a safe base prefix
                    base_path = file_path if file_path and os.path.exists(file_path) else f"temp_pdf_{datetime.utcnow().timestamp()}"
                else:
                    doc = fitz.open(file_path)
                    base_path = file_path

                # Process all pages (or up to a reasonable sanity limit like 50)
                # Requirement implies support for 50 page PDFs
                for i in range(min(50, len(doc))):
                    page = doc.load_page(i)
                    # Increase DPI to 300 for high-quality OCR on bad scans
                    pix = page.get_pixmap(dpi=300)
                    img_path = f"{base_path}_page_{i}.jpg"
                    pix.save(img_path)
                    image_paths.append(img_path)
                doc.close()
            except Exception as e:
                logger.error(f"PDF to Image conversion failed: {e}")
            return image_paths

        return await asyncio.to_thread(_convert)

    async def _extract_text_locally(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """
        Extracts text from a PDF using pypdf locally.
        Crucially, this extracts FORM FIELDS from editable PDFs.
        """
        def _read_pdf():
            text_content = []
            try:
                if file_content:
                    reader = PdfReader(io.BytesIO(file_content))
                else:
                    reader = PdfReader(file_path)
                
                # --- DIAGNOSTICS ---
                if reader.is_encrypted:
                    logger.warning(f"PDF is encrypted. Attempting to read anyway (might fail if password needed).")
                    try:
                        reader.decrypt("")
                    except:
                        pass
                
                # Check for XFA (Dynamic Forms)
                if "/AcroForm" in reader.trailer["/Root"] and "/XFA" in reader.trailer["/Root"]["/AcroForm"]:
                     logger.warning("PDF appears to contain XFA (Dynamic Form) data. Standard extraction might be limited.")
                     text_content.append("[WARNING: Document is an XFA Dynamic Form. Data might be hidden.]")
                # -------------------

                # 1. Extract Form Fields (Key for Editable PDFs)
                try:
                    fields = reader.get_form_text_fields()
                    if fields:
                        text_content.append("--- FORM FIELD DATA ---")
                        for key, value in fields.items():
                            if value:
                                text_content.append(f"{key}: {value}")
                        text_content.append("--- END FORM DATA ---\n")
                    else:
                        logger.info("No standard AcroForm fields found.")
                except Exception as e:
                    logger.warning(f"Failed to extract form fields: {e}")

                # 2. Extract Page Text
                for i, page in enumerate(reader.pages):
                    text_content.append(f"--- PAGE {i+1} ---")
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            # ══════════════════════════════════════════════════════════════════
                            # APPLY FRAGMENTED TEXT CLEANING - Fix for word-per-line PDFs
                            # ══════════════════════════════════════════════════════════════════
                            cleaned_page_text = clean_fragmented_text(page_text)
                            text_content.append(cleaned_page_text)
                            
                            # Log cleaning effectiveness for diagnostic purposes
                            if len(page_text) != len(cleaned_page_text):
                                logger.debug(f"Page {i+1} text cleaned: {len(page_text)} -> {len(cleaned_page_text)} chars")
                            # ══════════════════════════════════════════════════════════════════
                        else:
                            text_content.append("[EMPTY PAGE TEXT - LIKELY IMAGE OR XFA]")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {i+1}: {e}")
                        
            except Exception as e:
                logger.error(f"Local PDF reading failed: {e}")
                return ""
                
            return "\n".join(text_content)

        return await asyncio.to_thread(_read_pdf)

    async def _extract_text_from_docx(self, file_path: str, file_content: Optional[bytes] = None) -> str:
        """
        Extracts text from a DOCX file using python-docx.
        Extracts paragraphs (with heading styles), tables, headers, and footers.
        """
        def _read_docx():
            try:
                from docx import Document as DocxDocument
            except ImportError:
                logger.error("python-docx is not installed. Install with: pip install python-docx")
                return ""

            text_parts = []
            try:
                if file_content:
                    doc = DocxDocument(io.BytesIO(file_content))
                else:
                    doc = DocxDocument(file_path)

                # 1. Extract headers from all sections
                for i, section in enumerate(doc.sections):
                    try:
                        header = section.header
                        if header and header.paragraphs:
                            header_text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
                            if header_text.strip():
                                text_parts.append("--- DOCUMENT HEADER ---")
                                text_parts.append(header_text)
                                text_parts.append("--- END HEADER ---\n")
                    except Exception as e:
                        logger.warning(f"Failed to extract header from section {i}: {e}")

                # 2. Extract paragraphs with style info (helps LLM identify sections)
                for para in doc.paragraphs:
                    if para.text.strip():
                        style_name = para.style.name if para.style else ""
                        if "heading" in style_name.lower():
                            text_parts.append(f"\n## {para.text}")
                        else:
                            text_parts.append(para.text)

                # 3. Extract all tables (inventor info, addresses often live here)
                for table_idx, table in enumerate(doc.tables):
                    text_parts.append(f"\n--- TABLE {table_idx + 1} ---")
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        row_text = " | ".join(cells)
                        if any(cells):  # Skip empty rows
                            text_parts.append(row_text)
                    text_parts.append(f"--- END TABLE {table_idx + 1} ---\n")

                # 4. Extract footers
                for i, section in enumerate(doc.sections):
                    try:
                        footer = section.footer
                        if footer and footer.paragraphs:
                            footer_text = "\n".join(p.text for p in footer.paragraphs if p.text.strip())
                            if footer_text.strip():
                                text_parts.append("--- DOCUMENT FOOTER ---")
                                text_parts.append(footer_text)
                                text_parts.append("--- END FOOTER ---\n")
                    except Exception as e:
                        logger.warning(f"Failed to extract footer from section {i}: {e}")

            except Exception as e:
                logger.error(f"DOCX text extraction failed: {e}", exc_info=True)
                return ""

            full_text = "\n".join(text_parts)
            logger.info(f"DOCX extraction complete: {len(full_text)} chars, {len(text_parts)} text blocks")
            return full_text

        return await asyncio.to_thread(_read_docx)

    async def _analyze_docx_document(
        self,
        file_path: str,
        file_content: Optional[bytes] = None,
        progress_callback: Optional[Callable[[int, str], Awaitable[None]]] = None
    ) -> PatentApplicationMetadata:
        """
        Analyze a DOCX document by extracting text then using text-only analysis.
        DOCX files cannot be processed by Gemini natively, so we extract text first.
        """
        logger.info(f"--- ANALYZING DOCX DOCUMENT: {file_path} ---")

        if progress_callback:
            await progress_callback(10, "Extracting text from DOCX document...")

        # Extract all text from the DOCX
        text_content = await self._extract_text_from_docx(file_path, file_content)

        if not text_content or len(text_content.strip()) < 50:
            logger.error(f"DOCX text extraction yielded insufficient text: {len(text_content or '')} chars")
            raise ValueError("Could not extract meaningful text from the DOCX document.")

        logger.info(f"DOCX text extraction successful: {len(text_content)} characters")

        if progress_callback:
            await progress_callback(30, "Analyzing document content with AI...")

        # Use the same text-only analysis path used for PDFs with extractable text
        result = await self._analyze_text_only(text_content)

        if progress_callback:
            await progress_callback(100, "DOCX analysis complete")

        return result

    # --- DocuMind Extraction Pipeline Methods ---

    # async def extract_document(self, file_path: str) -> ExtractionResult:
    #     """
    #     Main entry point for document extraction.
    #     Decides whether to process directly or in chunks.
    #     """
    #     file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        
    #     # Get page count
    #     try:
    #         reader = PdfReader(file_path)
    #         page_count = len(reader.pages)
    #     except Exception:
    #         page_count = 0 # Fallback
            
    #     should_chunk = (
    #         file_size_mb > settings.LARGE_FILE_THRESHOLD_MB or 
    #         page_count > settings.LARGE_FILE_PAGE_THRESHOLD
    #     )
        
    #     with open(file_path, "rb") as f:
    #         file_bytes = f.read()

    #     if should_chunk:
    #         logger.info(f"Document requires chunking (Size: {file_size_mb:.2f}MB, Pages: {page_count})")
    #         return await self._extract_document_chunked(file_bytes, os.path.basename(file_path), page_count)
    #     else:
    #         logger.info(f"Processing document directly (Size: {file_size_mb:.2f}MB, Pages: {page_count})")
    #         # For direct extraction, we can reuse the file_path upload to save bandwidth 
    #         # if we hadn't already read bytes, but here we follow the pipeline logic
    #         return await self._extract_document_direct(file_path)

    # async def _extract_document_direct(self, file_path: str) -> ExtractionResult:
    #     """
    #     Extract text from a small document in a single API call using Native PDF support.
    #     """
    #     # Upload file to Gemini
    #     file_obj = await self.upload_file(file_path)
        
    #     extraction_prompt = """
    #     You are DocuMind, a High-Fidelity Document Digitization System.

    #     ## CORE PRINCIPLES
    #     1. **NO HALLUCINATION** - Never invent, assume, or fabricate any information not explicitly visible in the document.
    #     2. **NO SUMMARIZATION** - Extract the COMPLETE content of every page.
    #     3. **PRESERVE FIDELITY** - Maintain the original spelling, punctuation, formatting structure.

    #     ## OUTPUT FORMAT
    #     For each page, output in this exact format:
        
    #     --- PAGE [number] ---
        
    #     [Extract all visible text exactly as it appears, preserving structure]
        
    #     [Use annotations for non-text elements like: [Handwritten: ...], [Stamp: ...], [Table: ...]]

    #     After ALL pages, provide:
        
    #     === DOCUMENT EXTRACTION SUMMARY ===
    #     TOTAL PAGES: [count]
    #     OVERALL DOCUMENT CONFIDENCE: [High/Medium/Low]
    #     DOCUMENT QUALITY: [Excellent/Good/Fair/Poor]
    #     HANDWRITING DETECTED: [Yes/No]
    #     EXTRACTION NOTES: [Any important observations]
    #     """
        
    #     retries = settings.GEMINI_MAX_RETRIES
    #     for attempt in range(retries):
    #         try:
    #             # We use the raw client directly here to get text output, not JSON
    #             response = await asyncio.to_thread(
    #                 self.client.models.generate_content,
    #                 model=settings.GEMINI_MODEL,
    #                 contents=[file_obj, extraction_prompt],
    #                 config=types.GenerateContentConfig(
    #                     temperature=0.0,
    #                     max_output_tokens=65536
    #                 )
    #             )
                
    #             self._log_token_usage(response, "direct_extraction")
    #             extracted_text = response.text
    #             metadata_dict = self._parse_extraction_metadata(extracted_text, os.path.getsize(file_path))
                
    #             return ExtractionResult(
    #                 extracted_text=extracted_text,
    #                 metadata=ExtractionMetadata(**metadata_dict)
    #             )
            
    #         except ResourceExhausted as re_err:
    #             logger.warning(f"Gemini Rate Limit Exceeded during direct extraction: {re_err}")
    #             raise HTTPException(
    #                 status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
    #                 detail="AI Service is currently busy (Rate Limit Exceeded). Please try again in a moment."
    #             )
    #         except Exception as e:
    #             logger.warning(f"Direct extraction failed (attempt {attempt + 1}/{retries}): {e}")
    #             if attempt == retries - 1:
    #                 logger.error(f"Direct extraction failed after {retries} attempts: {e}")
    #                 raise e
                
    #             # Exponential backoff: 2s, 4s, 8s
    #             wait_time = (2 ** (attempt + 1))
    #             logger.info(f"Retrying in {wait_time} seconds...")
    #             await asyncio.sleep(wait_time)

    # async def _extract_document_chunked(self, file_bytes: bytes, filename: str, total_pages: int) -> ExtractionResult:
    #     """
    #     Extract text from a large document using parallel chunk processing.
    #     """
    #     # Split document into chunks
    #     chunks = self._chunk_pdf(file_bytes, settings.CHUNK_SIZE_PAGES)
    #     total_chunks = len(chunks)
        
    #     # Semaphore for concurrency control
    #     semaphore = asyncio.Semaphore(settings.MAX_CONCURRENT_EXTRACTIONS)

    #     async def extract_chunk(chunk_data: Tuple[bytes, int, int], chunk_index: int):
    #         chunk_bytes, start_page, end_page = chunk_data
    #         async with semaphore:
    #             return await self._extract_single_chunk(
    #                 chunk_bytes, chunk_index, total_chunks, start_page, end_page
    #             )

    #     # Process all chunks in parallel
    #     tasks = [extract_chunk(chunk_data, idx) for idx, chunk_data in enumerate(chunks)]
    #     results = await asyncio.gather(*tasks, return_exceptions=True)
        
    #     return self._aggregate_chunk_results(results, total_chunks, len(file_bytes))

    async def _analyze_document_chunked_structured(
        self,
        file_bytes: bytes,
        filename: str,
        total_pages: int,
        progress_callback: Optional[Callable[[int, str], Awaitable[None]]] = None
    ) -> PatentApplicationMetadata:
        """
        Analyzes a large document by splitting it into chunks and processing them in parallel
        to extract STRUCTURED metadata (Inventors, Title, etc.).
        """
        # 1. Split into chunks
        # Use a slightly larger chunk size for structured data to ensure context (e.g. 10 pages)
        chunk_size = 10
        chunks = self._chunk_pdf(file_bytes, chunk_size_pages=chunk_size)
        total_chunks = len(chunks)
        
        logger.info(f"Splitting {total_pages} pages into {total_chunks} chunks for Structured Analysis.")

        # 2. Parallel Processing
        semaphore = asyncio.Semaphore(settings.MAX_CONCURRENT_EXTRACTIONS)
        processed_count = 0

        async def process_chunk(chunk_data: Tuple[bytes, int, int], chunk_index: int):
            nonlocal processed_count
            chunk_bytes, start_page, end_page = chunk_data
            
            async with semaphore:
                logger.info(f"Starting Structured Analysis for Chunk {chunk_index + 1}/{total_chunks} (Pages {start_page}-{end_page})")
                try:
                    result = await self._extract_structured_chunk(
                        chunk_bytes, chunk_index, total_chunks, start_page, end_page
                    )
                    
                    processed_count += 1
                    if progress_callback:
                        # Map progress 20-90%
                        progress = 20 + int((processed_count / total_chunks) * 70)
                        await progress_callback(progress, f"Analyzed chunk {processed_count}/{total_chunks}")
                    
                    return result
                except Exception as e:
                    logger.error(f"Failed to analyze chunk {chunk_index}: {e}")
                    return None

        tasks = [process_chunk(c, i) for i, c in enumerate(chunks)]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 3. Aggregate Results
        # Filter out failed chunks
        valid_results = [r for r in results if r is not None and not isinstance(r, Exception)]
        
        return self._aggregate_structured_chunks(valid_results)

    async def _extract_structured_chunk(
        self, chunk_bytes: bytes, chunk_index: int, total_chunks: int,
        start_page: int, end_page: int, max_retries: int = 3
    ) -> Dict[str, Any]:
        """
        Extract structured metadata from a single PDF chunk.
        """
        chunk_prompt = f"""
        You are DocuMind. You are processing CHUNK {chunk_index+1} of {total_chunks} from a larger patent document.
        This chunk contains pages {start_page} to {end_page}.

        ## INSTRUCTIONS
        1. **Inventors (CRITICAL)**:
           - Scan EVERY PAGE in this chunk for "Inventor Information", "Legal Name", or similar tables.
           - Extract ALL inventors found.
           - If a list of inventors continues from a previous page, INCLUDE THEM.
        2. **Bibliographic Data**:
           - Look for Title, Application Number, Filing Date, Entity Status.
           - Note: These might only appear on the first page of the first chunk, but check anyway.
        3. **Applicant/Company**:
           - Look for "Applicant Information", "Assignee Information", or company details.
           - Extract organization name and address if found in this chunk.

        ## OUTPUT SCHEMA
        Return JSON with:
        - title (string/null)
        - application_number (string/null)
        - entity_status (string/null)
        - inventors (list of objects)
        - applicant (object with company information, or null if not found)
        """
        
        schema = {
            "title": "Title found (or null)",
            "application_number": "Application number (or null)",
            "entity_status": "Entity status (or null)",
            "inventors": [
                {
                    "name": "Full Name",
                    "first_name": "First name",
                    "middle_name": "Middle name",
                    "last_name": "Last name",
                    "city": "City",
                    "state": "State",
                    "country": "Country",
                    "street_address": "Street address / Mailing address"
                }
            ],
            "applicant": {
                "name": "Company/Applicant name",
                "street_address": "Street address",
                "city": "City",
                "state": "State",
                "zip_code": "Postal/ZIP code",
                "country": "Country"
            }
        }

        for attempt in range(max_retries):
            try:
                # Write chunk to temp file for upload
                temp_filename = f"temp_struct_chunk_{chunk_index}_{attempt}_{random.randint(1000,9999)}.pdf"
                with open(temp_filename, "wb") as f:
                    f.write(chunk_bytes)
                
                try:
                    file_obj = await self.upload_file(temp_filename)
                    
                    result = await self.generate_structured_content(
                        prompt=chunk_prompt,
                        file_obj=file_obj,
                        schema=schema
                    )
                    return result

                finally:
                    if os.path.exists(temp_filename):
                        try:
                            os.remove(temp_filename)
                        except:
                            pass

            except Exception as e:
                logger.warning(f"Chunk {chunk_index} failed attempt {attempt+1}: {e}")
                wait_time = (2 ** attempt) * 2
                await asyncio.sleep(wait_time)

        return {}

    def _aggregate_structured_chunks(self, results: List[Dict[str, Any]]) -> PatentApplicationMetadata:
        """
        Aggregates metadata from multiple chunks into a single result.
        """
        final_metadata = {
            "title": None,
            "application_number": None,
            "entity_status": None,
            "inventors": [],
            "applicant": None
        }
        
        extracted_inventors = []
        
        for res in results:
            if not res: continue
            
            # 1. Metadata (First valid wins)
            if not final_metadata["title"] and res.get("title"):
                final_metadata["title"] = res["title"]
            if not final_metadata["application_number"] and res.get("application_number"):
                final_metadata["application_number"] = res["application_number"]
            if not final_metadata["entity_status"] and res.get("entity_status"):
                final_metadata["entity_status"] = res["entity_status"]
            if not final_metadata["applicant"] and res.get("applicant"):
                final_metadata["applicant"] = res["applicant"]
            
            # 2. Inventors (Merge and Deduplicate)
            # FIXED: Two inventors with same name but different addresses are DIFFERENT people
            if res.get("inventors"):
                for new_inv in res["inventors"]:
                    is_duplicate = False
                    
                    # Build full name for comparison
                    new_name = new_inv.get("name", "").strip().lower()
                    if not new_name and (new_inv.get("first_name") or new_inv.get("last_name")):
                        parts = [p for p in [new_inv.get("first_name"), new_inv.get("middle_name"), new_inv.get("last_name")] if p]
                        new_name = " ".join(parts).strip().lower()

                    for existing in extracted_inventors:
                        existing_name = existing.get("name", "").strip().lower()
                        if not existing_name:
                            parts = [p for p in [existing.get("first_name"), existing.get("middle_name"), existing.get("last_name")] if p]
                            existing_name = " ".join(parts).strip().lower()
                        
                        # If names don't match, not a duplicate
                        if not new_name or not existing_name or new_name != existing_name:
                            continue
                        
                        # Names match - check if addresses also match (true duplicate)
                        def normalize(val):
                            return (val or "").strip().lower()
                        
                        same_street = normalize(existing.get("street_address")) == normalize(new_inv.get("street_address")) and normalize(new_inv.get("street_address")) != ""
                        same_city = normalize(existing.get("city")) == normalize(new_inv.get("city")) and normalize(new_inv.get("city")) != ""
                        same_state = normalize(existing.get("state")) == normalize(new_inv.get("state")) and normalize(new_inv.get("state")) != ""
                        same_zip = normalize(existing.get("zip_code")) == normalize(new_inv.get("zip_code")) and normalize(new_inv.get("zip_code")) != ""
                        
                        # Consider duplicate if: same name AND (same street OR (same city AND state) OR same zip)
                        address_match = same_street or (same_city and same_state) or same_zip
                        
                        # If both have no address info, treat same name as duplicate
                        both_no_address = (
                            not existing.get("street_address") and not existing.get("city") and
                            not existing.get("state") and not existing.get("zip_code") and
                            not new_inv.get("street_address") and not new_inv.get("city") and
                            not new_inv.get("state") and not new_inv.get("zip_code")
                        )
                        
                        if address_match or both_no_address:
                            is_duplicate = True
                            # Merge fields if existing is empty but new has data
                            if not existing.get("city") and new_inv.get("city"):
                                existing["city"] = new_inv["city"]
                            if not existing.get("state") and new_inv.get("state"):
                                existing["state"] = new_inv["state"]
                            if not existing.get("street_address") and new_inv.get("street_address"):
                                existing["street_address"] = new_inv["street_address"]
                            if not existing.get("zip_code") and new_inv.get("zip_code"):
                                existing["zip_code"] = new_inv["zip_code"]
                            if not existing.get("country") and new_inv.get("country"):
                                existing["country"] = new_inv["country"]
                            break
                    
                    if not is_duplicate:
                        extracted_inventors.append(new_inv)
        
        final_metadata["inventors"] = extracted_inventors
        
        # Post-processing: Name splitting
        for inventor in final_metadata["inventors"]:
            if inventor.get("name") and not inventor.get("last_name"):
                parts = inventor["name"].split()
                if len(parts) >= 2:
                    inventor["first_name"] = parts[0]
                    inventor["last_name"] = parts[-1]
                    if len(parts) > 2:
                        inventor["middle_name"] = " ".join(parts[1:-1])
                elif len(parts) == 1:
                    inventor["first_name"] = parts[0]

        return PatentApplicationMetadata(**final_metadata)

    def _chunk_pdf(self, pdf_bytes: bytes, chunk_size_pages: int = 5) -> List[Tuple[bytes, int, int]]:
        """
        Split a PDF into chunks of specified page count.
        """
        reader = PdfReader(io.BytesIO(pdf_bytes))
        total_pages = len(reader.pages)
        chunks = []

        for start_idx in range(0, total_pages, chunk_size_pages):
            end_idx = min(start_idx + chunk_size_pages, total_pages)

            writer = PdfWriter()
            for page_idx in range(start_idx, end_idx):
                writer.add_page(reader.pages[page_idx])

            chunk_buffer = io.BytesIO()
            writer.write(chunk_buffer)
            chunk_bytes = chunk_buffer.getvalue()

            chunks.append((
                chunk_bytes,
                start_idx + 1,      # start_page (1-indexed)
                end_idx             # end_page (1-indexed)
            ))

        return chunks

    async def _extract_single_chunk(
        self, chunk_bytes: bytes, chunk_index: int, total_chunks: int, 
        start_page: int, end_page: int, max_retries: int = 3
    ) -> dict:
        """
        Extract text from a single chunk with retry logic.
        """
        chunk_prompt = f"""
        You are DocuMind. You are processing CHUNK {chunk_index+1} of {total_chunks} from a larger document.
        This chunk contains pages {start_page} to {end_page}.
        
        ## CORE PRINCIPLES
        1. **NO HALLUCINATION**
        2. **NO SUMMARIZATION**
        3. **PRESERVE FIDELITY**

        ## OUTPUT FORMAT
        For each page in this chunk, use this format:
        
        --- PAGE [actual page number, starting at {start_page}] ---
        
        [Full extraction content]
        
        [Page Confidence: High/Medium/Low]

        ## CHUNK SUMMARY
        After extracting all pages in this chunk, provide:
        
        === CHUNK {chunk_index+1} EXTRACTION SUMMARY ===
        PAGES IN CHUNK: {start_page}-{end_page}
        CHUNK CONFIDENCE: [High/Medium/Low]
        """

        for attempt in range(max_retries):
            try:
                # We need to upload the chunk as a file to Gemini
                # Ideally we would save to temp file, but for now lets try passing bytes if SDK supports
                # SDK might require path, so let's write to a temp file
                temp_filename = f"temp_chunk_{chunk_index}_{attempt}.pdf"
                with open(temp_filename, "wb") as f:
                    f.write(chunk_bytes)
                
                try:
                    file_obj = await self.upload_file(temp_filename)
                    
                    response = await asyncio.to_thread(
                        self.client.models.generate_content,
                        model=settings.GEMINI_MODEL,
                        contents=[file_obj, chunk_prompt],
                        config=types.GenerateContentConfig(
                            temperature=0.0,
                            max_output_tokens=65536
                        )
                    )

                    self._log_token_usage(response, f"chunk_extraction_{chunk_index}")
                    
                    return {
                        "chunk_index": chunk_index,
                        "extracted_text": response.text,
                        "success": True
                    }
                finally:
                    if os.path.exists(temp_filename):
                        os.remove(temp_filename)

            except Exception as e:
                wait_time = (2 ** attempt) * 2
                await asyncio.sleep(wait_time)

        return {
            "chunk_index": chunk_index,
            "extracted_text": f"[EXTRACTION FAILED FOR PAGES {start_page}-{end_page}]",
            "success": False
        }

    # def _aggregate_chunk_results(self, results: List[Any], total_chunks: int, file_size: int) -> ExtractionResult:
    #     """
    #     Combine extraction results from multiple chunks.
    #     """
    #     successful_results = [r for r in results if isinstance(r, dict) and r.get("success")]
    #     successful_results.sort(key=lambda x: x["chunk_index"])

    #     combined_text_parts = []
    #     for result in successful_results:
    #         combined_text_parts.append(result["extracted_text"])

    #     combined_text = "\n\n".join(combined_text_parts)
        
    #     failed_count = total_chunks - len(successful_results)
        
    #     metadata = ExtractionMetadata(
    #         page_count=0, # Would need to parse from text or pass through
    #         overall_confidence=ConfidenceLevel.LOW if failed_count > 0 else ConfidenceLevel.HIGH,
    #         is_chunked=True,
    #         chunk_count=total_chunks,
    #         successful_chunks=len(successful_results),
    #         failed_chunks=failed_count,
    #         file_size_bytes=file_size
    #     )
        
    #     return ExtractionResult(extracted_text=combined_text, metadata=metadata)

    # def _parse_extraction_metadata(self, extracted_text: str, file_size: int) -> dict:
    #     """
    #     Parse metadata from LLM extraction output.
    #     """
    #     metadata = {
    #         "page_count": 0,
    #         "overall_confidence": "medium",
    #         "document_quality": "good",
    #         "has_handwriting": False,
    #         "extraction_notes": None,
    #         "file_size_bytes": file_size,
    #         "uncertain_count": 0,
    #         "illegible_count": 0
    #     }

    #     # Count pages from markers
    #     page_markers = re.findall(r'--- PAGE (\d+) ---', extracted_text)
    #     if page_markers:
    #         metadata["page_count"] = len(page_markers)

    #     # Extract overall confidence
    #     confidence_match = re.search(
    #         r'OVERALL DOCUMENT CONFIDENCE:\s*(High|Medium|Low)',
    #         extracted_text,
    #         re.IGNORECASE
    #     )
    #     if confidence_match:
    #         metadata["overall_confidence"] = confidence_match.group(1).lower()

    #     # Extract document quality
    #     quality_match = re.search(
    #         r'DOCUMENT QUALITY:\s*(Excellent|Good|Fair|Poor)',
    #         extracted_text,
    #         re.IGNORECASE
    #     )
    #     if quality_match:
    #         metadata["document_quality"] = quality_match.group(1).lower()

    #     # Check for handwriting
    #     handwriting_match = re.search(
    #         r'HANDWRITING DETECTED:\s*(Yes|No)',
    #         extracted_text,
    #         re.IGNORECASE
    #     )
    #     if handwriting_match:
    #         metadata["has_handwriting"] = handwriting_match.group(1).lower() == "yes"

    #     # Also check for handwriting annotations
    #     if re.search(r'\[Handwritten:', extracted_text):
    #         metadata["has_handwriting"] = True

    #     return metadata
    async def analyze_office_action(
        self,
        file_path: str,
        file_content: Optional[bytes] = None,
        progress_callback: Optional[Callable[[int, str], Awaitable[None]]] = None,
        extract_claim_text: bool = False
    ) -> Dict[str, Any]:
        """
        Analyzes a Patent Office Action PDF.
        Extracts structured data: Header, Claims Status, Rejections, Objections, etc.

        Args:
            file_path: Path to the Office Action PDF
            file_content: Optional bytes content if already loaded
            progress_callback: Optional callback for progress updates
            extract_claim_text: If True, attempts to extract full claim text from the document
        """
        from app.models.office_action import OfficeActionExtractedData

        logger.info(f"--- ANALYZING OFFICE ACTION: {file_path} (extract_claim_text={extract_claim_text}) ---")

        # Upload file for multimodal analysis
        if file_content:
            upload_source = io.BytesIO(file_content)
        else:
            upload_source = file_path
            
        try:
            if progress_callback:
                await progress_callback(10, "Uploading Office Action for AI Analysis...")
                
            file_obj = await self.upload_file(upload_source)

            # Build conditional claims text extraction instruction
            claims_text_instruction = ""
            if extract_claim_text:
                claims_text_instruction = """
            10. **FULL CLAIMS TEXT EXTRACTION**:
                - For EACH claim, extract the complete claim text if present in the Office Action
                - Extract the preamble (introductory phrase like "A method for...", "A system comprising...")
                - Extract the body (the claim elements/limitations)
                - Note: Claims may be reproduced in the Office Action for rejection context
                - If claims are not present in the document, leave claim_text, claim_preamble, claim_body as null
"""

            prompt = """
            Analyze the provided Patent Office Action PDF.
            Extract structured information with high precision.
            IMPORTANT: Return null for any field that is not present in the document. Do NOT return "N/A" or empty strings for missing data.

            ################################################################################
            ################################################################################
            ##                                                                            ##
            ##  MANDATORY: FIND ALL §103 OBVIOUSNESS REJECTIONS                          ##
            ##                                                                            ##
            ##  THIS IS THE #1 PRIORITY. §103 IS THE MOST COMMON REJECTION TYPE.         ##
            ##  FAILURE TO DETECT §103 REJECTIONS IS A CRITICAL ERROR.                   ##
            ##                                                                            ##
            ################################################################################
            ################################################################################

            ## HOW TO IDENTIFY §103 vs §102:

            §102 (ANTICIPATION) - uses "ANTICIPATED BY" with ONE reference:
            > "Claims X are rejected under 35 U.S.C. 102 as being ANTICIPATED BY Smith."

            §103 (OBVIOUSNESS) - uses "UNPATENTABLE OVER...IN VIEW OF" with MULTIPLE references:
            > "Claims X are rejected under 35 U.S.C. 103 as being UNPATENTABLE OVER Malladi IN VIEW OF Tu."
            > "Claims X are rejected under 35 U.S.C. 103 as being UNPATENTABLE OVER Malladi IN VIEW OF Dollinger."
            > "Claim X is rejected under 35 U.S.C. 103 as being UNPATENTABLE OVER Malladi and Dollinger IN VIEW OF Tu."

            ## STEP-BY-STEP §103 DETECTION PROCESS:

            STEP 1: Search the ENTIRE document for text containing "103" or "unpatentable over"
            STEP 2: For each occurrence, check if it says "rejected under 35 U.S.C. 103"
            STEP 3: Extract the PRIMARY reference (appears after "over" or "unpatentable over")
            STEP 4: Extract SECONDARY reference(s) (appear after "in view of" or "combined with")
            STEP 5: Extract the affected claim numbers from the rejection heading
            STEP 6: Each UNIQUE combination of references = SEPARATE rejection entry

            ## REQUIRED §103 PATTERNS TO SEARCH FOR:

            Pattern 1: "rejected under 35 U.S.C. 103 as being unpatentable over"
            Pattern 2: "rejected under 35 U.S.C. § 103"
            Pattern 3: "rejected under pre-AIA 35 U.S.C. 103"
            Pattern 4: "unpatentable over [name] in view of [name]"
            Pattern 5: "obvious over [name] in view of [name]"
            Pattern 6: "It would have been obvious to one of ordinary skill"
            Pattern 7: "would have been motivated to combine"

            ## EXAMPLE FROM REAL OFFICE ACTIONS:

            If you see this text:
            > "Claims 3, 4, 6, 7, 17, 18, 22 and 24 are rejected under 35 U.S.C. 103 as being
            > unpatentable over Malladi (US Pub. No. 2017/0060574) in view of Tu (US Pub. No. 2006/0250248)."

            You MUST create a rejection entry with:
            - rejection_type: "103"
            - statutory_basis: "35 U.S.C. 103"
            - affected_claims: ["3", "4", "6", "7", "17", "18", "22", "24"]
            - Primary reference: Malladi
            - Secondary reference: Tu

            If you see ANOTHER §103 rejection with DIFFERENT references:
            > "Claims 8, 9, 11, 12, 13, 14, 15 are rejected under 35 U.S.C. 103 as being
            > unpatentable over Malladi in view of Dollinger (NPL)."

            You MUST create a SEPARATE rejection entry with:
            - rejection_type: "103"
            - affected_claims: ["8", "9", "11", "12", "13", "14", "15"]
            - Primary reference: Malladi
            - Secondary reference: Dollinger

            ## TRIPLE COMBINATIONS:

            If you see:
            > "Claim 10 is rejected under 35 U.S.C. 103 as being unpatentable over Malladi and Dollinger,
            > and further in view of Tu."

            Create a rejection with:
            - affected_claims: ["10"]
            - Primary: Malladi
            - Secondary: [Dollinger, Tu]

            ## COMMON §103 DETECTION MISTAKES TO AVOID:

            1. DO NOT skip §103 rejections just because you found §102 rejections
            2. DO NOT combine multiple §103 rejections into one entry - each reference combo is separate
            3. DO NOT confuse "anticipated by" (§102) with "unpatentable over" (§103)
            4. DO NOT stop reading after §112 rejections - §103 rejections often appear later in the document
            5. DO NOT ignore NPL (Non-Patent Literature) as secondary references

            ################################################################################
            ## INSTRUCTIONS
            ################################################################################

            1. **HEADER INFO**: Extract comprehensive header information including:
               - Application Number (look for "Serial No.", "Application No.", "Appl. No.")
               - Filing Date
               - Office Action Date (Mailing Date)
               - Office Action Type: "Non-Final", "Final", "Advisory", "Restriction", "Ex Parte Quayle"
               - Examiner Name, Art Unit
               - Attorney Docket Number, Confirmation Number
               - First Named Inventor (usually visible in header or first page)
               - Applicant Name/Entity (company or individual applicant)
               - **Title of Invention** - MANDATORY EXTRACTION:
                 * This field MUST be populated if the document contains a title
                 * Look for the longest capitalized phrase describing the invention
                 * Search ALL of these locations - do NOT stop at first failure:
                   1. Line starting with "Title:" (the text after "Title:")
                   2. Line starting with "Re:" (the text after "Re:")
                   3. Line starting with "For:" (the text after "For:")
                   4. Near "Application No." or "Serial No." header
                   5. First page top section in ALL CAPS
                   6. FILING DATA/APPLICATION INFORMATION section
                   7. Any line containing "SYSTEMS AND METHODS" or "METHOD AND APPARATUS"
                 * Title is usually a long phrase in ALL CAPS like:
                   "SYSTEMS AND METHODS FOR AUTOMATIC TRANSFERRING OF..."
                   "METHOD FOR DATA PROCESSING IN DISTRIBUTED NETWORKS"
               - Customer Number (in correspondence address section)
               - Examiner's Phone Number and Email (usually at end of document)
               - Examiner designation (Primary Examiner vs Assistant Examiner)

            2. **FOREIGN PRIORITY DATA**:
               - Look for "FOREIGN PRIORITY DATA" or "35 U.S.C. 119" priority claims
               - Extract: country, application number, filing date for each priority claim

            3. **PARENT APPLICATION DATA (CONTINUITY)**:
               - Look for "DOMESTIC PRIORITY DATA", "CONTINUITY DATA", or "RELATED U.S. APPLICATION DATA"
               - Extract parent application numbers, filing dates, and relationship types
               - Relationship types: "continuation of", "continuation-in-part of", "divisional of"

            4. **CLAIMS STATUS** - CRITICAL for accurate classification:
               - List ALL claims mentioned in the document
               - Determine status: Rejected, Allowed, Objected to, Cancelled, Withdrawn, Pending
               - **INDEPENDENT vs DEPENDENT CLASSIFICATION**:
                 * A claim is DEPENDENT if it references another claim. Look for these patterns:
                   - "The [noun] of claim [N]" (e.g., "The method of claim 1")
                   - "The [noun] according to claim [N]"
                   - "The [noun] as recited in claim [N]"
                   - "A [noun] as in claim [N]"
                   - "Claim [N], wherein..."
                   - "The [noun] of any of claims [N-M]"
                   - "The [noun] of any one of claims [N], [M], or [O]"
                 * A claim is INDEPENDENT if it does NOT reference any other claim
                 * Common independent claim starts: "A method...", "A system...", "An apparatus...", "A computer-readable medium..."
               - **USE §112(d) REJECTIONS TO DETERMINE DEPENDENCIES**:
                 * §112(d) rejections explicitly state claim dependency relationships
                 * Example: "Claim 20 is rejected under 35 U.S.C. 112(d) as being of improper dependent form... Claim 20 depends on claim 19"
                 * If a §112(d) rejection says "Claim X depends on claim Y", then:
                   - Claim X is DEPENDENT (not Independent!)
                   - parent_claim = "Y"
                 * Always cross-check claim classifications with §112(d) rejections
               - **PARENT CLAIM** - For EVERY dependent claim, you MUST identify the ACTUAL parent:
                 * Extract the claim number from the dependency phrase in the REJECTION TEXT
                 * The examiner quotes the claim language: read "the [noun] of claim N" carefully
                 * "The method of claim 8" → parent_claim = "8"
                 * "The system of claim 1" → parent_claim = "1"
                 * "The centralized system of claim 2" → parent_claim = "2" (NOT claim 1!)
                 * "The machine of claim 12" → parent_claim = "12" (NOT claim 8!)
                 
                 ⚠️ CRITICAL MISTAKE TO AVOID: DO NOT default all dependent claims in a group
                 to the same independent claim! Dependent claims typically form a CHAIN:
                   - Claim 2 depends on claim 1 ("The system of claim 1")
                   - Claim 3 depends on claim 2 ("The system of claim 2")  ← NOT claim 1!
                   - Claim 4 depends on claim 3 ("The system of claim 3")  ← NOT claim 1!
                   - Claim 5 depends on claim 1 ("The system of claim 1")  ← Back to claim 1
                   - Claim 6 depends on claim 5 ("The system of claim 5")  ← NOT claim 1!
                   - Claim 7 depends on claim 6 ("The system of claim 6")  ← NOT claim 1!
                 
                 Each claim has its OWN parent — read each "of claim N" phrase individually.
                 
                 * For multiple dependency, use the first claim: "any of claims 1-5" → parent_claim = "1"
                 * **Check §112(d) rejections** - they explicitly name parent claims
                 * **VERIFY**: After assigning parents, confirm that no two claims in a chain
                   all point to the same independent claim unless the text explicitly says so.

            5. **REJECTIONS (Critical)** - YOU MUST CAPTURE ALL REJECTION TYPES:
               - Extract EACH rejection block separately
               - READ THE ENTIRE DOCUMENT - rejections appear on multiple pages

               **REJECTION CATEGORIES TO FIND:**

               A) §112 Rejections (usually appear first):
                  - 112(a): written description/enablement
                  - 112(b): indefiniteness
                  - 112(d): improper dependent claim form

               B) §102 Rejections (anticipation):
                  - Look for: "anticipated by [single reference]"
                  - Only ONE reference cited

               C) §103 Rejections (obviousness) - MOST COMMON, DON'T MISS THESE:
                  - Look for: "unpatentable over [Ref A] in view of [Ref B]"
                  - MULTIPLE references combined
                  - Create SEPARATE entries for each unique reference combination

               **FOR EACH §103 REJECTION YOU FIND:**
               1. Set rejection_type = "103"
               2. Set rejection_type_normalized = "103"
               3. Set statutory_basis = "35 U.S.C. 103"
               4. Extract affected_claims from the rejection header
               5. Extract primary reference (after "unpatentable over" or "obvious over")
               6. Extract secondary reference(s) (after "in view of")
               7. Extract motivation_to_combine (after "It would have been obvious...")

               - **REJECTION TYPE DETECTION** - Identify precise statutory basis:
                 * 35 U.S.C. 101 - Patent eligibility (abstract idea, laws of nature)
                 * 35 U.S.C. 102(a)(1) - Prior art; uses "anticipated by" (AIA)
                 * 35 U.S.C. 102(a)(2) - Prior art; uses "anticipated by" (AIA)
                 * 35 U.S.C. 102(a), 102(b), 102(e), 102(g) - Pre-AIA novelty; "anticipated by"
                 * 35 U.S.C. 103 - Obviousness; uses "unpatentable over...in view of" (AIA) - VERY COMMON
                 * 35 U.S.C. 103(a) - Obviousness (pre-AIA); "unpatentable over...in view of"
                 * 35 U.S.C. 112(a) - Written description, enablement (AIA)
                 * 35 U.S.C. 112(b) - Definiteness/indefiniteness (AIA)
                 * 35 U.S.C. 112(d) - Improper dependent claim (ALSO reveals claim dependencies!)
                 * 35 U.S.C. 112(f) - Means-plus-function interpretation issues
                 * 35 U.S.C. 112 first/second paragraph - Pre-AIA equivalents
                 * Double Patenting - Statutory or Obviousness-type (ODP)

               - **GROUPING §103 REJECTIONS**: Create SEPARATE rejection entries for different claim/reference combinations:
                 * Example: Claims 3,4,6,7 rejected over Malladi+Tu is ONE rejection
                 * Example: Claims 8,9,11-15 rejected over Malladi+Dollinger is a SEPARATE rejection
                 * Example: Claim 10 rejected over Malladi+Dollinger+Tu is a THIRD rejection
                 * Each unique combination of references = new rejection entry
               - Set is_aia to true for AIA citations (102(a)(1), 102(a)(2), 112(a), 112(b)), false for pre-AIA
               - Set rejection_type_normalized to the specific subsection (e.g., "102(a)(1)", "112(b)", "103")
               - List affected claim numbers
               - Extract Examiner's Reasoning verbatim or summarized

            6. **PRIOR ART REFERENCES** - Extract ALL references, including NPL:
               - Extract ALL prior art references into the all_references list
               - This includes references from §102 AND §103 rejections
               - Assign unique reference_id (ref_1, ref_2, ref_3, etc.)
               - **REFERENCE TYPES:**
                 * "US Patent" - e.g., "US 9,999,999 B2"
                 * "US Publication" - e.g., "US 2017/0060574 A1", "US Pub. No. 2006/0250248"
                 * "Foreign Patent" - e.g., "EP 1234567 A1"
                 * "NPL" - Non-Patent Literature (academic papers, articles)
                   - Example: Dollinger, "A Framework for Efficient Execution on GPU..."
                   - For NPL, use author name as short_name
               - Extract for each reference:
                 * reference_type: one of the above types
                 * identifier: Full citation
                 * short_name: Examiner's shorthand (e.g., "Malladi", "Tu", "Dollinger")
                 * inventor_author: First named inventor or author
                 * title: Document title if available
                 * date: Publication or issue date
                 * relevant_sections: Cited paragraphs/columns
               - Track which rejection indices use each reference in used_in_rejection_indices
               - **IMPORTANT**: If a reference appears in multiple rejections, list all rejection indices

            7. **SECTION 103 PRIOR ART COMBINATIONS** - CRITICAL FOR §103 REJECTIONS:
               - For EACH 35 U.S.C. 103 rejection, identify the combination structure:
                 * Primary reference: The FIRST reference mentioned (the base reference)
                 * Secondary reference(s): References combined with primary
                   - Look for: "in view of", "combined with", "in combination with", "and"
                   - Example: "obvious over Malladi in view of Tu" → Primary: Malladi, Secondary: Tu
                   - Example: "obvious over A in view of B and C" → Primary: A, Secondary: [B, C]
                 * Motivation to combine: Examiner's reasoning for combination
                   - Look for: "It would have been obvious...", "One of ordinary skill would have been motivated..."
                   - Look for: "to provide...", "in order to...", "for the purpose of..."
               - Create SEPARATE prior_art_combinations entries for different claim groups within the same 103 rejection

            8. **OBJECTIONS**:
               - Identify objections to Specification, Drawings, Claims, Title, Abstract
               - Categorize objection_type: "claim", "specification", "drawings", "title", "abstract", "formality"
               - Extract reason and required corrective action

            9. **OTHER STATEMENTS**:
                - Look for "Allowable Subject Matter" indications
                - Interview summaries
                - Examiner suggestions
                - Response Deadline

            10. **SUPERVISOR INFORMATION** (Usually at end of document):
                - Look for "Supervisory Patent Examiner" in signature block
                - Extract supervisor name (e.g., "BORIS GORNEY")
                - Extract supervisor phone (usually 10-digit number like "5712705626")
                - Pattern: "[NAME] can be reached on [PHONE]"
                - Pattern: "/[NAME]/ Supervisory Patent Examiner, Art Unit XXXX"

            11. **FAX NUMBER**:
                - Look for "fax phone number" or "fax number"
                - Pattern: "fax phone number for the organization is XXX-XXX-XXXX"
                - Usually 571-273-XXXX for USPTO

            12. **AMENDED CLAIMS** (Critical for tracking prosecution history):
                - Look for "Status of Claims" section near beginning
                - Pattern: "Claims X, Y, Z are currently amended"
                - Pattern: "Claims X-Y have been amended"
                - Extract ALL claim numbers that were amended in this response

            13. **IDS (Information Disclosure Statement) DATES**:
                - Look for "Information Disclosure Statement(s)" acknowledgment
                - Pattern: "Paper No(s)/Mail Date X and Y"
                - Pattern: "IDS filed on X has been considered"
                - Extract all submission dates

            14. **EXTENSION PERIOD DETAILS**:
                - Look for "A shortened statutory period for reply"
                - Extract statutory_period_months (usually 3)
                - Look for "maximum statutory period"
                - Extract max_extension_months (usually 6)

            15. **APPLICANT'S ARGUMENTS STATUS**:
                - Look for "Applicant's arguments regarding claims..."
                - Status types: "moot", "persuasive", "not persuasive"
                - Pattern: "arguments ... are fully considered but moot in view of..."
                - Extract which claims the arguments addressed
                - Extract the reason (e.g., "moot in view of the new grounds of rejection")
            """ + claims_text_instruction + """

            ################################################################################
            ## MANDATORY FINAL VERIFICATION - DO THIS BEFORE OUTPUTTING JSON
            ################################################################################

            Before generating output, you MUST verify:

            VERIFICATION 1 - §103 REJECTIONS:
            □ Search the document for "unpatentable over" - did you find ALL instances?
            □ Search for "in view of" - each occurrence with "103" is a potential §103 rejection
            □ Count your §103 rejections - most Office Actions have 2-5 §103 rejections
            □ If you found 0 §103 rejections, GO BACK AND SEARCH AGAIN - you likely missed them
            □ Did you create SEPARATE entries for different reference combinations?

            VERIFICATION 2 - TITLE OF INVENTION:
            □ Is title_of_invention populated? Search for "Title:", "Re:", "For:" lines
            □ Look for ALL CAPS text on the first page describing the invention
            □ The title should NOT be null unless truly absent from the document

            VERIFICATION 3 - CLAIM DEPENDENCIES:
            □ For any claim with §112(d) rejection, mark as Dependent with correct parent_claim
            □ §112(d) text explicitly states "Claim X depends on claim Y"

            ################################################################################
            ## OUTPUT SCHEMA
            ################################################################################

            Return JSON matching the following structure. Note the examples showing:
            - Multiple §103 rejections with different reference combinations (using "in view of" language)
            - Claims with proper Independent/Dependent classification and parent_claim values
            - Use §112(d) rejections to determine claim dependencies

            {
                "header": {
                    "application_number": "string or null",
                    "filing_date": "string or null",
                    "office_action_date": "string or null",
                    "office_action_type": "string or null",
                    "examiner_name": "string or null",
                    "art_unit": "string or null",
                    "attorney_docket_number": "string or null",
                    "confirmation_number": "string or null",
                    "response_deadline": "string or null",
                    "first_named_inventor": "string or null",
                    "applicant_name": "string or null",
                    "title_of_invention": "REQUIRED - Extract from Title:, Re:, For: lines or ALL CAPS text on first page",
                    "customer_number": "string or null",
                    "examiner_phone": "string or null",
                    "examiner_email": "string or null",
                    "examiner_type": "string or null",
                    "foreign_priority": [
                        { "country": "...", "application_number": "...", "filing_date": "..." }
                    ],
                    "parent_applications": [
                        { "parent_application_number": "...", "parent_filing_date": "...", "relationship_type": "...", "status": "..." }
                    ]
                },
                "claims_status": [
                    { "claim_number": "1", "status": "Rejected", "dependency_type": "Independent", "parent_claim": null },
                    { "claim_number": "2", "status": "Rejected", "dependency_type": "Dependent", "parent_claim": "1" },
                    { "claim_number": "8", "status": "Rejected", "dependency_type": "Independent", "parent_claim": null },
                    { "claim_number": "9", "status": "Rejected", "dependency_type": "Dependent", "parent_claim": "8" },
                    { "claim_number": "12", "status": "Rejected", "dependency_type": "Dependent", "parent_claim": "8" }
                ],
                "rejections": [
                    {
                        "rejection_type": "112",
                        "rejection_type_normalized": "112(b)",
                        "statutory_basis": "35 U.S.C. 112(b)",
                        "is_aia": true,
                        "affected_claims": ["1", "2"],
                        "examiner_reasoning": "Claims are indefinite because...",
                        "cited_prior_art": [],
                        "prior_art_combinations": []
                    },
                    {
                        "rejection_type": "103",
                        "rejection_type_normalized": "103",
                        "statutory_basis": "35 U.S.C. 103",
                        "is_aia": true,
                        "affected_claims": ["1", "2", "3", "4"],
                        "examiner_reasoning": "Claims 1-4 are rejected as obvious over Malladi in view of Tu...",
                        "cited_prior_art": [
                            {
                                "reference_id": "ref_1",
                                "reference_type": "US Publication",
                                "identifier": "US 2017/0060574 A1",
                                "short_name": "Malladi",
                                "inventor_author": "Malladi et al.",
                                "relevant_sections": "[0190], [0193]",
                                "relevant_claims": ["1", "2", "3", "4"]
                            },
                            {
                                "reference_id": "ref_2",
                                "reference_type": "US Publication",
                                "identifier": "US 2018/0123456 A1",
                                "short_name": "Tu",
                                "inventor_author": "Tu et al.",
                                "relevant_sections": "[0045]",
                                "relevant_claims": ["1", "2", "3", "4"]
                            }
                        ],
                        "prior_art_combinations": [
                            {
                                "primary_reference_id": "ref_1",
                                "secondary_reference_ids": ["ref_2"],
                                "motivation_to_combine": "It would have been obvious to combine Malladi with Tu to provide improved performance..."
                            }
                        ]
                    },
                    {
                        "rejection_type": "103",
                        "rejection_type_normalized": "103",
                        "statutory_basis": "35 U.S.C. 103",
                        "is_aia": true,
                        "affected_claims": ["8", "9", "10", "11"],
                        "examiner_reasoning": "Claims 8-11 are rejected as obvious over Malladi in view of Dollinger...",
                        "cited_prior_art": [
                            {
                                "reference_id": "ref_1",
                                "reference_type": "US Publication",
                                "identifier": "US 2017/0060574 A1",
                                "short_name": "Malladi"
                            },
                            {
                                "reference_id": "ref_3",
                                "reference_type": "NPL",
                                "identifier": "Dollinger, 'A Framework for Efficient Execution...'",
                                "short_name": "Dollinger"
                            }
                        ],
                        "prior_art_combinations": [
                            {
                                "primary_reference_id": "ref_1",
                                "secondary_reference_ids": ["ref_3"],
                                "motivation_to_combine": "One of ordinary skill would have been motivated to..."
                            }
                        ]
                    },
                    {
                        "rejection_type": "103",
                        "rejection_type_normalized": "103",
                        "statutory_basis": "35 U.S.C. 103",
                        "is_aia": true,
                        "affected_claims": ["10"],
                        "examiner_reasoning": "Claim 10 is rejected as obvious over Malladi and Dollinger in view of Tu...",
                        "cited_prior_art": [
                            { "reference_id": "ref_1", "short_name": "Malladi" },
                            { "reference_id": "ref_3", "short_name": "Dollinger" },
                            { "reference_id": "ref_2", "short_name": "Tu" }
                        ],
                        "prior_art_combinations": [
                            {
                                "primary_reference_id": "ref_1",
                                "secondary_reference_ids": ["ref_3", "ref_2"],
                                "motivation_to_combine": "..."
                            }
                        ]
                    }
                ],
                "objections": [
                    {
                        "objected_item": "Claim 1",
                        "objection_type": "claim",
                        "reason": "...",
                        "corrective_action": "..."
                    }
                ],
                "other_statements": [
                    { "statement_type": "Allowable Subject Matter", "content": "..." }
                ],
                "all_references": [
                    {
                        "reference_id": "ref_1",
                        "reference_type": "US Publication",
                        "identifier": "US 2017/0060574 A1",
                        "short_name": "Malladi",
                        "inventor_author": "Malladi et al.",
                        "title": "Method for...",
                        "date": "2017-03-02",
                        "relevant_sections": "[0190], [0193], [0200]",
                        "used_in_rejection_indices": [1, 2, 3]
                    },
                    {
                        "reference_id": "ref_2",
                        "reference_type": "US Publication",
                        "identifier": "US 2006/0250248",
                        "short_name": "Tu",
                        "used_in_rejection_indices": [1, 3]
                    },
                    {
                        "reference_id": "ref_3",
                        "reference_type": "NPL",
                        "identifier": "Dollinger, 'A Framework for Efficient Execution on GPU and CPU+GPU Systems'",
                        "short_name": "Dollinger",
                        "inventor_author": "Dollinger",
                        "used_in_rejection_indices": [2, 3]
                    }
                ]
            }
            """
            
            # Build claims_status schema - conditionally include claim text fields
            claims_status_schema = {
                "claim_number": "string",
                "status": "string",
                "dependency_type": "string",
                "parent_claim": "string or null"
            }
            if extract_claim_text:
                claims_status_schema["claim_text"] = "string or null"
                claims_status_schema["claim_preamble"] = "string or null"
                claims_status_schema["claim_body"] = "string or null"

            schema = {
                "header": {
                    "application_number": "string or null",
                    "filing_date": "string or null",
                    "office_action_date": "string or null",
                    "office_action_type": "string or null",
                    "examiner_name": "string or null",
                    "art_unit": "string or null",
                    "attorney_docket_number": "string or null",
                    "confirmation_number": "string or null",
                    "response_deadline": "string or null",
                    "first_named_inventor": "string or null",
                    "applicant_name": "string or null",
                    "title_of_invention": "string or null",
                    "customer_number": "string or null",
                    "examiner_phone": "string or null",
                    "examiner_email": "string or null",
                    "examiner_type": "string or null",
                    "supervisor_name": "string or null",
                    "supervisor_phone": "string or null",
                    "fax_number": "string or null",
                    "statutory_period_months": "integer or null",
                    "max_extension_months": "integer or null",
                    "foreign_priority": [
                        {
                            "country": "string or null",
                            "application_number": "string or null",
                            "filing_date": "string or null",
                            "certified_copy_attached": "boolean or null"
                        }
                    ],
                    "parent_applications": [
                        {
                            "parent_application_number": "string or null",
                            "parent_filing_date": "string or null",
                            "relationship_type": "string or null",
                            "status": "string or null",
                            "patent_number": "string or null"
                        }
                    ]
                },
                "claims_status": [claims_status_schema],
                "rejections": [
                    {
                        "rejection_type": "string",
                        "rejection_type_normalized": "string or null",
                        "statutory_basis": "string or null",
                        "is_aia": "boolean or null",
                        "affected_claims": ["string"],
                        "examiner_reasoning": "string",
                        "cited_prior_art": [
                            {
                                "reference_id": "string or null",
                                "reference_type": "string",
                                "identifier": "string",
                                "short_name": "string or null",
                                "inventor_author": "string or null",
                                "title": "string or null",
                                "date": "string or null",
                                "relevant_sections": "string or null",
                                "relevant_claims": ["string"]
                            }
                        ],
                        "prior_art_combinations": [
                            {
                                "primary_reference_id": "string",
                                "secondary_reference_ids": ["string"],
                                "motivation_to_combine": "string or null",
                                "teaching_suggestion_motivation": "string or null",
                                "affected_claim_elements": ["string"]
                            }
                        ]
                    }
                ],
                "objections": [
                    {
                        "objected_item": "string",
                        "objection_type": "string or null",
                        "reason": "string",
                        "corrective_action": "string or null"
                    }
                ],
                "other_statements": [
                    {
                        "statement_type": "string",
                        "content": "string"
                    }
                ],
                "all_references": [
                    {
                        "reference_id": "string",
                        "reference_type": "string",
                        "identifier": "string",
                        "short_name": "string or null",
                        "inventor_author": "string or null",
                        "title": "string or null",
                        "date": "string or null",
                        "relevant_sections": "string or null",
                        "used_in_rejection_indices": ["integer"]
                    }
                ],
                "amended_claims": ["string"],
                "ids_submissions": [
                    {
                        "submission_date": "string",
                        "paper_number": "string or null",
                        "was_considered": "boolean"
                    }
                ],
                "applicant_arguments": [
                    {
                        "status": "string",
                        "affected_claims": ["string"],
                        "reason": "string or null"
                    }
                ]
            }

            if progress_callback:
                await progress_callback(30, "AI Analyzing Document Structure...")

            result = await self.generate_structured_content(
                prompt=prompt,
                file_obj=file_obj,
                schema=schema
            )
            
            if progress_callback:
                await progress_callback(90, "Finalizing Extraction...")

            # ── POST-PROCESSING: Apply validation and fixes ──────────────────────
            try:
                # Extract PDF text for post-processing validation
                if file_content:
                    pdf_text = await self._extract_text_locally(file_path, file_content)
                else:
                    pdf_text = await self._extract_text_locally(file_path)
                
                # Apply post-processing fixes
                result = _post_process_office_action(result, pdf_text, logger)
                logger.info("Post-processing validation completed successfully")
                
            except Exception as post_error:
                logger.warning(f"Post-processing failed but continuing with original result: {post_error}")
                # Continue with original result if post-processing fails

            return result

        except Exception as e:
            logger.error(f"Office Action Analysis Failed: {e}")
            raise e

llm_service = LLMService()